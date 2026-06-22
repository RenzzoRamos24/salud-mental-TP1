"""
Documento Word completo de Sami para presentar al asesor de tesis.
Lenguaje sencillo, sin tecnicismos, como si se le explicara a un alumno de
15 años. Cubre: qué es el sistema, objetivo, alcance, roles, consentimiento
de padres, cómo funciona el diario, palabras clave por ítem, tabla
días→puntos, ciclo de 14 días, cálculo PHQ-A/GAD-7/DSM-5, qué ve cada rol.

Salida: docs/Sami_para_asesor.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paleta de colores ───────────────────────────────────────────────
VERDE = RGBColor(0x05, 0x96, 0x69)
VERDE_OSC = RGBColor(0x04, 0x6A, 0x49)
GRIS_OSCURO = RGBColor(0x0A, 0x0A, 0x0A)
GRIS_MEDIO = RGBColor(0x5A, 0x5A, 0x5A)
AMBAR = RGBColor(0xB8, 0x85, 0x0A)
ROJO = RGBColor(0xB0, 0x21, 0x21)


# ── Helpers de formato ──────────────────────────────────────────────
def h1(doc, texto):
    p = doc.add_heading(texto, level=1)
    for run in p.runs:
        run.font.color.rgb = VERDE
        run.font.size = Pt(20)


def h2(doc, texto):
    p = doc.add_heading(texto, level=2)
    for run in p.runs:
        run.font.color.rgb = VERDE_OSC
        run.font.size = Pt(15)


def h3(doc, texto):
    p = doc.add_heading(texto, level=3)
    for run in p.runs:
        run.font.color.rgb = GRIS_OSCURO
        run.font.size = Pt(12)


def parrafo(doc, texto, negrita=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(size)
    if negrita:
        run.bold = True
    if italic:
        run.italic = True
    return p


def parrafo_mixto(doc, partes, size=11):
    """partes = [(texto, {'bold':True, 'italic':False, 'color':RGB}), ...]"""
    p = doc.add_paragraph()
    for texto, fmt in partes:
        run = p.add_run(texto)
        run.font.size = Pt(size)
        if fmt.get("bold"):
            run.bold = True
        if fmt.get("italic"):
            run.italic = True
        if fmt.get("color"):
            run.font.color.rgb = fmt["color"]
    return p


def cita(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(texto)
    run.italic = True
    run.font.color.rgb = GRIS_MEDIO
    run.font.size = Pt(11)


def bullet(doc, texto, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(texto)
    run.font.size = Pt(size)


def numerado(doc, texto, size=11):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(texto)
    run.font.size = Pt(size)


def alerta(doc, texto, color=AMBAR):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("⚠  " + texto)
    run.font.size = Pt(11)
    run.font.color.rgb = color
    run.bold = True


def tabla(doc, encabezados, filas, ancho_col=None):
    t = doc.add_table(rows=len(filas) + 1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, enc in enumerate(encabezados):
        hdr[i].text = enc
    for c in hdr:
        for pr in c.paragraphs:
            for run in pr.runs:
                run.bold = True
    for i, fila in enumerate(filas, 1):
        for j, val in enumerate(fila):
            t.rows[i].cells[j].text = val
    return t


def salto(doc):
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


# ── PORTADA ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\nSami")
run.bold = True
run.font.size = Pt(44)
run.font.color.rgb = VERDE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Diario emocional escolar con monitoreo asistido por IA")
run.italic = True
run.font.size = Pt(16)
run.font.color.rgb = GRIS_MEDIO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "\n\nDocumento explicativo del sistema\n"
    "para la presentación al asesor de tesis\n\n"
    "Universidad Peruana de Ciencias Aplicadas — UPC"
)
run.font.size = Pt(12)
run.font.color.rgb = GRIS_OSCURO

salto(doc)


# ════════════════════════════════════════════════════════════════════
# 1. ¿QUÉ ES SAMI?
# ════════════════════════════════════════════════════════════════════
h1(doc, "1. ¿Qué es Sami?")

parrafo(
    doc,
    "Sami es un sistema digital pensado para los estudiantes del colegio. "
    "Es como un cuaderno de notas privado donde cada alumno puede contar "
    "cómo se siente. Lo que escribe no se lo cuenta a nadie: se queda en "
    "el sistema. Pero atrás de ese cuaderno hay una inteligencia artificial "
    "que lee con cuidado lo que el alumno escribió y va anotando señales "
    "sobre cómo está su salud emocional.",
)

parrafo(
    doc,
    "Si esas señales se repiten muchos días seguidos, el sistema se lo "
    "avisa a la psicóloga del colegio para que pueda acompañar al alumno. "
    "Si en cambio el alumno está bien o solo tiene un mal día, no pasa "
    "nada: se queda como un registro personal.",
)

parrafo(
    doc,
    "Sami no diagnostica. No dice “este alumno tiene depresión”. Lo que "
    "hace es contar señales y avisar. La que decide qué hacer siempre es "
    "la psicóloga.",
    italic=True,
)


# ════════════════════════════════════════════════════════════════════
# 2. OBJETIVO
# ════════════════════════════════════════════════════════════════════
h1(doc, "2. Objetivo del sistema")

parrafo(doc, "El objetivo principal es:")
bullet(
    doc,
    "Detectar de forma temprana señales de ansiedad, depresión y otras "
    "dificultades emocionales en estudiantes escolares.",
)
bullet(
    doc,
    "Darle al área de psicología del colegio información ordenada y "
    "objetiva para poder intervenir a tiempo.",
)
bullet(
    doc,
    "Generar un espacio seguro y privado para que el alumno se "
    "exprese sin sentirse juzgado.",
)
bullet(
    doc,
    "Apoyar el cumplimiento de la Ley 29733 (protección de datos "
    "personales) y el lineamiento del MINEDU sobre acompañamiento "
    "socioemocional en colegios.",
)


# ════════════════════════════════════════════════════════════════════
# 3. ALCANCE
# ════════════════════════════════════════════════════════════════════
h1(doc, "3. Alcance: qué hace y qué NO hace")

h3(doc, "Lo que Sami sí hace")
bullet(doc, "Recoge lo que el alumno escribe en su diario.")
bullet(doc, "Detecta señales emocionales con inteligencia artificial.")
bullet(
    doc,
    "Cuenta cuántos días aparece cada señal a lo largo de un ciclo de 14 días.",
)
bullet(
    doc,
    "Calcula puntajes oficiales de las escalas PHQ-A (depresión) y "
    "GAD-7 (ansiedad).",
)
bullet(
    doc,
    "Marca cuando un patrón observado coincide con los criterios "
    "mínimos del DSM-5.",
)
bullet(
    doc,
    "Avisa a la psicóloga si hay una posible crisis (por ejemplo, "
    "ideas de hacerse daño).",
)
bullet(doc, "Permite a la psicóloga agendar y registrar citas.")

h3(doc, "Lo que Sami NO hace")
bullet(doc, "No diagnostica enfermedades.")
bullet(doc, "No reemplaza a la psicóloga ni al tratamiento profesional.")
bullet(
    doc,
    "No le muestra al alumno etiquetas como “tienes depresión” ni le "
    "asigna puntajes en su cara.",
)
bullet(
    doc,
    "No comparte el contenido literal del diario con los padres ni con "
    "profesores. Solo la psicóloga tiene acceso, y aún ella ve solo el "
    "análisis, no el texto crudo.",
)
bullet(
    doc,
    "No se usa para sancionar, calificar ni reportar al alumno en su "
    "expediente escolar.",
)


# ════════════════════════════════════════════════════════════════════
# 4. ROLES DEL SISTEMA
# ════════════════════════════════════════════════════════════════════
h1(doc, "4. ¿Quiénes usan Sami? Los tres roles")

parrafo(
    doc,
    "El sistema tiene tres tipos de usuario, cada uno con permisos "
    "distintos. Nadie ve lo que no le corresponde.",
)

h2(doc, "4.1 Estudiante")
bullet(doc, "Es el dueño del diario.")
bullet(doc, "Escribe entradas cuando quiere (no es obligatorio diariamente).")
bullet(
    doc,
    "Ve consejos de autocuidado y, si lo pide, líneas de ayuda "
    "(Línea 113, etc.).",
)
bullet(
    doc,
    "NO ve ningún diagnóstico ni puntaje. Solo recibe acompañamiento "
    "en lenguaje cálido.",
)
bullet(
    doc,
    "Si pasa por una crisis, ve en pantalla las líneas de emergencia y "
    "el aviso de que la psicóloga ya fue notificada.",
)

h2(doc, "4.2 Psicóloga del colegio")
bullet(doc, "Tiene acceso a un panel de monitoreo con todos sus estudiantes.")
bullet(
    doc,
    "Ve, por cada alumno, sus puntajes PHQ-A y GAD-7 del ciclo actual, "
    "los ítems del DSM-5 que se están cumpliendo, y si hay posibles riesgos.",
)
bullet(doc, "Puede entrar al historial detallado de un estudiante.")
bullet(doc, "Puede agendar, completar o cancelar citas.")
bullet(
    doc,
    "Puede dejar notas clínicas privadas (que el alumno nunca ve).",
)
bullet(
    doc,
    "Puede enviarle al alumno un mensaje motivacional corto, sin "
    "diagnóstico.",
)

h2(doc, "4.3 Administrador del sistema")
bullet(doc, "Gestiona usuarios: registra, activa, desactiva.")
bullet(doc, "Configura las preguntas y los parámetros del modelo.")
bullet(doc, "Hace respaldos automáticos de la base de datos.")
bullet(
    doc,
    "Audita los accesos: cada vez que alguien entra al sistema queda registrado.",
)
bullet(doc, "NO ve el contenido del diario de ningún alumno.")


# ════════════════════════════════════════════════════════════════════
# 5. CONSENTIMIENTO DE LOS PADRES
# ════════════════════════════════════════════════════════════════════
h1(doc, "5. Consentimiento de los padres (muy importante)")

parrafo(
    doc,
    "Como los estudiantes son menores de edad, antes de que un alumno "
    "pueda usar Sami por primera vez el sistema le muestra un documento "
    "de consentimiento informado que deben firmar los padres o el "
    "apoderado.",
)

h3(doc, "Qué dice ese consentimiento, en palabras simples")
bullet(
    doc,
    "Que Sami es una herramienta para apoyar el bienestar emocional "
    "del alumno.",
)
bullet(
    doc,
    "Que lo que el alumno escriba en su diario será leído únicamente "
    "por un sistema de inteligencia artificial y por la psicóloga del "
    "colegio. Nadie más.",
)
bullet(
    doc,
    "Que los datos están protegidos según la Ley N° 29733 (protección "
    "de datos personales en el Perú).",
)
bullet(
    doc,
    "Que el sistema NO diagnostica. Solo detecta señales y avisa a la "
    "psicóloga.",
)
bullet(
    doc,
    "Que si el sistema detecta una posible crisis (por ejemplo, ideas "
    "de hacerse daño), la psicóloga será avisada de inmediato y se "
    "contactará a los padres siguiendo el protocolo del colegio.",
)
bullet(
    doc,
    "Que los padres pueden, en cualquier momento, pedir que su hijo "
    "deje de usar Sami y que sus datos sean eliminados.",
)

h3(doc, "Cómo se firma")
parrafo(
    doc,
    "El consentimiento queda registrado dentro del sistema con la fecha "
    "y la versión del documento. Si la versión cambia (por ejemplo, si "
    "se agregan funciones nuevas), se les pide a los padres firmar la "
    "versión actualizada.",
)

alerta(
    doc,
    "Sin firma del padre o apoderado, el alumno no puede ingresar al "
    "diario. El sistema lo bloquea automáticamente.",
)


# ════════════════════════════════════════════════════════════════════
# 6. CÓMO FUNCIONA, EN PALABRAS SENCILLAS
# ════════════════════════════════════════════════════════════════════
h1(doc, "6. Cómo funciona Sami, paso a paso")

parrafo(
    doc,
    "Imaginemos a María, una alumna de 4° de secundaria.",
    italic=True,
)

h3(doc, "Paso 1 — María se registra")
parrafo(
    doc,
    "María entra al sistema con su correo del colegio, crea una "
    "contraseña y ve el documento de consentimiento. Sus padres lo "
    "firman desde su propio celular. Recién ahí Sami se le abre.",
)

h3(doc, "Paso 2 — María abre el diario por primera vez")
parrafo(
    doc,
    "Sami le hace una pregunta suave, como por ejemplo:",
)
cita(
    doc,
    "“Hola María. Cuéntame, ¿cómo te fue hoy? ¿Hubo algo que te hiciera "
    "sentir bien o algo que te molestara?”",
)
parrafo(
    doc,
    "María escribe lo que quiera, con sus propias palabras. Por ejemplo:",
)
cita(
    doc,
    "“Hoy me sentí triste. No tengo ganas de nada y dormí muy poco. "
    "Siento que estoy fallándole a mis papás.”",
)

h3(doc, "Paso 3 — Sami lee lo que María escribió")
parrafo(
    doc,
    "En segundo plano, sin que María lo vea, el sistema hace tres cosas:",
)
numerado(
    doc,
    "Busca palabras clave. En el ejemplo, encuentra: “triste”, “no "
    "tengo ganas”, “dormí muy poco”, “fallándole”. Cada una activa "
    "un ítem distinto.",
)
numerado(
    doc,
    "Con un modelo de inteligencia artificial llamado BETO (es BERT en "
    "español), Sami estima qué tan fuerte aparece cada señal: si fue "
    "algo leve, fuerte, o si parece presente “casi todo el tiempo”.",
)
numerado(
    doc,
    "Guarda esa información asociada al día en que María escribió.",
)

h3(doc, "Paso 4 — María cierra el diario, sigue su vida normal")
parrafo(
    doc,
    "El alumno NUNCA ve un puntaje ni un diagnóstico. Solo ve un mensaje "
    "amigable de Sami, como:",
)
cita(
    doc,
    "“Gracias por contarme, María. Recuerda que dormir bien también es "
    "cuidarte. Mañana te espero por aquí.”",
)

h3(doc, "Paso 5 — Si María vuelve a escribir al día siguiente")
parrafo(
    doc,
    "El sistema vuelve a leer, vuelve a contar palabras clave y vuelve "
    "a registrar señales. Día a día se va construyendo un patrón. Por "
    "eso es importante que escriba seguido, aunque sea poco.",
)

h3(doc, "Paso 6 — Si Sami detecta una crisis")
parrafo(
    doc,
    "Si en lo que María escribe aparecen señales graves (por ejemplo "
    "“ya no quiero estar”, “quiero desaparecer”), el sistema NO espera. "
    "Manda una alerta inmediata a la psicóloga, le muestra a María en "
    "pantalla las líneas de emergencia (Línea 113, SALUDLINE 106, "
    "Teléfono de la Esperanza) y deja registrado un evento de crisis.",
)

salto(doc)


# ════════════════════════════════════════════════════════════════════
# 7. PHQ-A, GAD-7, DSM-5 — TRES NOMBRES, EXPLICADOS FÁCIL
# ════════════════════════════════════════════════════════════════════
h1(doc, "7. PHQ-A, GAD-7 y DSM-5 explicados fácil")

h2(doc, "7.1 PHQ-A")
parrafo(
    doc,
    "PHQ-A son las siglas de Patient Health Questionnaire para "
    "adolescentes. Es un cuestionario clínico oficial, usado en "
    "hospitales y colegios de todo el mundo, que sirve para medir "
    "síntomas de depresión.",
)
parrafo(doc, "Tiene 9 preguntas, una por cada síntoma de depresión:")
bullet(doc, "¿Sigues disfrutando lo que antes te gustaba?")
bullet(doc, "¿Te has sentido triste o sin esperanza?")
bullet(doc, "¿Has dormido bien?")
bullet(doc, "¿Estás cansado/a todo el tiempo?")
bullet(doc, "¿Has cambiado en tu apetito?")
bullet(doc, "¿Te sientes culpable o inútil?")
bullet(doc, "¿Te cuesta concentrarte?")
bullet(doc, "¿Te mueves más lento o estás muy inquieto/a?")
bullet(doc, "¿Has tenido pensamientos de hacerte daño?")
parrafo(
    doc,
    "Cada una se contesta con 4 opciones: Nunca, Algunos días, Más de la "
    "mitad de los días, Casi todos los días. A esas opciones se les "
    "pone un puntaje de 0 a 3. Sumando los 9 ítems sale un total entre "
    "0 y 27.",
)

h2(doc, "7.2 GAD-7")
parrafo(
    doc,
    "GAD-7 son las siglas de Generalized Anxiety Disorder. También es "
    "un cuestionario clínico oficial, pero sirve para medir ansiedad. "
    "Tiene 7 preguntas:",
)
bullet(doc, "¿Te has sentido nervioso/a o con ansiedad?")
bullet(doc, "¿Te cuesta dejar de preocuparte?")
bullet(doc, "¿Te preocupas demasiado por muchas cosas?")
bullet(doc, "¿Te cuesta relajarte?")
bullet(doc, "¿Estás inquieto/a, como si no pudieras quedarte quieto?")
bullet(doc, "¿Estás irritable, de mal humor?")
bullet(doc, "¿Sientes que algo malo va a pasar?")
parrafo(
    doc,
    "Mismo sistema de respuesta: Nunca a Casi todos los días. Cada "
    "pregunta vale de 0 a 3. Total: entre 0 y 21.",
)

h2(doc, "7.3 DSM-5")
parrafo(
    doc,
    "DSM-5 es el manual oficial donde están definidos todos los "
    "trastornos de salud mental. Es como el libro de reglas de la "
    "psiquiatría. NO es un cuestionario. Lo que tiene son listas de "
    "criterios: para decir que alguien tiene depresión, qué tiene que "
    "estar pasando; para decir que tiene ansiedad, qué tiene que estar "
    "pasando.",
)
parrafo(
    doc,
    "En Sami, el DSM-5 se usa para dos cosas: (1) cada ítem del PHQ-A "
    "y del GAD-7 está etiquetado con su criterio DSM-5 correspondiente "
    "(por ejemplo, el ítem 2 del PHQ-A es el criterio “Estado de ánimo "
    "deprimido”), y (2) cuando varios criterios se cumplen, Sami "
    "marca esa combinación como “posible riesgo” según el DSM-5.",
)
parrafo(
    doc,
    "Sami nunca dice “el alumno tiene depresión”. Dice “el patrón "
    "observado coincide con los criterios mínimos del DSM-5 para "
    "Episodio Depresivo Mayor”. La diferencia es importante: el "
    "diagnóstico siempre lo pone una persona, no la máquina.",
    italic=True,
)


# ════════════════════════════════════════════════════════════════════
# 8. PALABRAS CLAVE — LISTA COMPLETA
# ════════════════════════════════════════════════════════════════════
h1(doc, "8. Las palabras que Sami busca")

parrafo(
    doc,
    "Cada ítem del PHQ-A y del GAD-7 tiene su lista de palabras clave. "
    "Si alguna aparece en lo que escribió el alumno, ese ítem queda "
    "“activado” para ese día. Después la inteligencia artificial decide "
    "qué tan fuerte está esa señal (0, 1, 2 ó 3 en la escala Likert).",
)

h2(doc, "8.1 PHQ-A — Palabras que indican depresión")
filas_phq = [
    (
        "Ítem 1 — Pérdida de interés (anhedonia)",
        '“no disfruto”, “nada me da alegría”, “no me gusta nada”, '
        '“perdí el interés”, “ya nada me importa”, “sin ganas”',
    ),
    (
        "Ítem 2 — Ánimo deprimido",
        '“triste”, “deprim…”, “vacío”, “sin esperanza”, '
        '“desesperanza”, “decaíd…”',
    ),
    (
        "Ítem 3 — Sueño",
        '“no puedo dormir”, “insomnio”, “duermo mucho”, '
        '“duermo todo el día”, “me despierto”, “no concilio”, '
        '“trasnoch…”',
    ),
    (
        "Ítem 4 — Fatiga",
        '“cansad…”, “agotad…”, “sin energía”, “fatiga”, '
        '“no tengo fuerzas”, “exhaust…”',
    ),
    (
        "Ítem 5 — Apetito",
        '“no tengo hambre”, “como mucho”, “atracón”, '
        '“perdí el apetito”, “comer de más”, “sin apetito”',
    ),
    (
        "Ítem 6 — Culpa / inutilidad",
        '“inútil”, “fracasad…”, “soy un fracaso”, “no sirvo”, '
        '“le fallé”, “culpa”, “no valgo nada”, “decepción”',
    ),
    (
        "Ítem 7 — Concentración",
        '“no me concentro”, “me distra…”, “no puedo concentrarme”, '
        '“se me va la cabeza”, “no retengo”, “leer y no entiendo”',
    ),
    (
        "Ítem 8 — Lentitud o agitación",
        '“muy lento”, “me muevo lento”, “no puedo quedarme quieto”, '
        '“inquiet…”, “agitad…”, “acelerad…”',
    ),
    (
        "Ítem 9 — Pensamientos de daño (CRÍTICO)",
        '“hacerme daño”, “mejor muerto”, “no quiero vivir”, '
        '“suicid…”, “desaparecer”, “no despertar”, “terminar con todo”, '
        '“autolesión”, “cortar”',
    ),
]
tabla(doc, ["Ítem", "Palabras que activan el ítem"], filas_phq)

alerta(
    doc,
    "El ítem 9 es crítico: si aparece cualquiera de esas palabras y la "
    "IA confirma señal, se dispara protocolo de emergencia inmediato.",
    color=ROJO,
)

h2(doc, "8.2 GAD-7 — Palabras que indican ansiedad")
filas_gad = [
    (
        "Ítem 1 — Nervios / ansiedad",
        '“nervios”, “ansi…”, “tens…”, “con los nervios”, '
        '“nerviosa”, “nerviosismo”',
    ),
    (
        "Ítem 2 — Preocupación incontrolable",
        '“no puedo dejar de preocuparme”, “no controlo”, '
        '“no paro de pensar”, “sobrepensar”, “rumiar”, “le doy vueltas”',
    ),
    (
        "Ítem 3 — Preocuparse por muchas cosas",
        '“me preocupo por todo”, “todo me preocupa”, '
        '“preocupación constante”, “preocup…”',
    ),
    (
        "Ítem 4 — Tensión / no relajarse",
        '“no me puedo relajar”, “no logro relajarme”, “tens…”, '
        '“rígido”, “contracturad…”',
    ),
    (
        "Ítem 5 — Inquietud",
        '“inquiet…”, “no me quedo quiet…”, “tengo que moverme”, '
        '“al límite”',
    ),
    (
        "Ítem 6 — Irritabilidad",
        '“irrit…”, “me molesto”, “me enojo rápido”, '
        '“todo me molesta”, “explot…”, “malhumorad…”',
    ),
    (
        "Ítem 7 — Sensación de catástrofe",
        '“algo malo va a pasar”, “siento que va a pasar algo”, '
        '“miedo”, “pánico”, “catástrofe”, “presentimiento”',
    ),
]
tabla(doc, ["Ítem", "Palabras que activan el ítem"], filas_gad)


# ════════════════════════════════════════════════════════════════════
# 9. TABLA DÍAS → PUNTOS
# ════════════════════════════════════════════════════════════════════
h1(doc, "9. La tabla mágica: días con síntoma → puntos")

parrafo(
    doc,
    "Esta es la parte más importante para defender ante el asesor. El "
    "PHQ-A y el GAD-7 originales se contestan con 4 opciones: Nunca, "
    "Algunos días, Más de la mitad de los días, Casi todos los días. "
    "Pero el alumno en Sami no contesta opciones; escribe libremente. "
    "Entonces, ¿cómo sabemos qué puntaje ponerle a cada ítem?",
)

parrafo(
    doc,
    "Sencillo: contamos en cuántos días distintos apareció ese síntoma "
    "durante el ciclo de 14 días, y convertimos ese número en el "
    "puntaje oficial usando esta tabla:",
    negrita=True,
)

filas_tabla = [
    ("0 días", "0 puntos", "Nunca"),
    ("1 a 7 días", "1 punto", "Algunos días"),
    ("8 a 11 días", "2 puntos", "Más de la mitad de los días"),
    ("12 a 14 días", "3 puntos", "Casi todos los días"),
]
tabla(
    doc,
    ["Días con el síntoma", "Puntos", "Frase oficial"],
    filas_tabla,
)

h3(doc, "Por qué esos cortes y no otros")
parrafo(
    doc,
    "Esta tabla NO la inventamos nosotros. Sale directo de la traducción "
    "literal de las palabras del cuestionario PHQ original (Johnson, 2002):",
)
bullet(
    doc,
    "“Nunca” significa, literalmente, ningún día. Por eso 0 días = 0 puntos.",
)
bullet(
    doc,
    "“Más de la mitad de los días” es una expresión matemática: la "
    "mitad de 14 es 7, así que “más de la mitad” es 8 o más. Por eso 8 a "
    "11 días = 2 puntos.",
)
bullet(
    doc,
    "“Casi todos los días” es claramente la franja alta, de 12 a 14.",
)
bullet(
    doc,
    "“Algunos días” es el rango que queda en el medio: 1 a 7 días.",
)
parrafo(
    doc,
    "Esta justificación es importante porque es defendible ante "
    "cualquier jurado: no estamos inventando reglas nuevas, estamos "
    "traduciendo las frases del cuestionario oficial a números, "
    "usando lo que esas frases dicen.",
    italic=True,
)


# ════════════════════════════════════════════════════════════════════
# 10. EL CICLO DE 14 DÍAS
# ════════════════════════════════════════════════════════════════════
h1(doc, "10. El ciclo de 14 días: cómo funciona el tiempo en Sami")

h2(doc, "10.1 ¿Por qué 14 días?")
parrafo(
    doc,
    "Porque el PHQ-A y el GAD-7 originales evalúan los “últimos 14 días” "
    "(2 semanas) del paciente. Así está validado clínicamente. Si "
    "usáramos 7 días sería poco para detectar un cuadro depresivo, y si "
    "usáramos 30 sería demasiado y se diluiría la señal.",
)

h2(doc, "10.2 ¿Cuándo empieza el ciclo de un alumno?")
parrafo(
    doc,
    "Cada alumno tiene su propio reloj. El Día 1 del ciclo es el día en "
    "que escribe por primera vez en el diario. A partir de ahí se "
    "cuentan 14 días corridos.",
)

h2(doc, "10.3 ¿Qué pasa durante esos 14 días?")
bullet(
    doc,
    "El alumno escribe cuando quiere. No es obligatorio diario, pero "
    "mientras más escriba, más confiable es el reporte.",
)
bullet(
    doc,
    "Cada entrada se analiza al instante: BETO lee, las palabras clave "
    "activan ítems, se guarda el día y la fuerza de cada señal.",
)
bullet(
    doc,
    "Si en algún momento aparece una señal de crisis, NO se espera al "
    "día 14: se alerta a la psicóloga al instante.",
)

h2(doc, "10.4 ¿Qué pasa al día 14?")
parrafo(
    doc,
    "El sistema cierra el ciclo automáticamente y genera el reporte "
    "clínico:",
)
numerado(doc, "Cuenta, ítem por ítem, en cuántos días distintos apareció.")
numerado(doc, "Convierte ese conteo en puntaje 0-3 usando la tabla Johnson.")
numerado(doc, "Suma todos los ítems PHQ-A → puntaje total 0-27.")
numerado(doc, "Suma todos los ítems GAD-7 → puntaje total 0-21.")
numerado(doc, "Calcula el nivel de severidad según los puntos de corte oficiales.")
numerado(doc, "Verifica si se cumplen criterios mínimos del DSM-5.")
numerado(doc, "Marca el ciclo como “cerrado” y guarda el reporte para la psicóloga.")

h2(doc, "10.5 ¿Y al día 15?")
parrafo(
    doc,
    "Empieza automáticamente el Ciclo 2 del alumno. El sistema no "
    "espera nada. Si la psicóloga tiene una cita programada con el "
    "alumno para el día 19, esa cita ocurre EN PARALELO al Ciclo 2 que "
    "ya está corriendo. Las cosas que el alumno escribe entre el día 15 "
    "y el día 19 cuentan para el Ciclo 2, no se anexan al Ciclo 1.",
)

h2(doc, "10.6 Excepción: cita marcada como crisis")
parrafo(
    doc,
    "Si la psicóloga atiende al alumno en una situación de emergencia "
    "y marca esa cita como “atención de crisis”, el ciclo en curso se "
    "cierra ese mismo día con reporte parcial. Al día siguiente arranca "
    "un ciclo nuevo. La razón: después de una crisis y su intervención, "
    "el contexto emocional cambia, y es mejor empezar de cero que "
    "arrastrar días bajo la sombra de la crisis.",
)


# ════════════════════════════════════════════════════════════════════
# 11. CÓMO SE CALCULA EL RESULTADO FINAL
# ════════════════════════════════════════════════════════════════════
h1(doc, "11. Cómo se calcula el resultado final")

h2(doc, "11.1 Ejemplo concreto con María")
parrafo(
    doc,
    "Imaginemos que en su ciclo de 14 días, María escribió en el diario "
    "en 10 días distintos (cobertura alta). El sistema detectó:",
)
filas_maria = [
    ("phq9_1", "Pérdida de interés", "PHQ-A", "8 días", "2 puntos", "Más de la mitad"),
    ("phq9_2", "Ánimo deprimido", "PHQ-A", "9 días", "2 puntos", "Más de la mitad"),
    ("phq9_3", "Problemas de sueño", "PHQ-A", "12 días", "3 puntos", "Casi todos"),
    ("phq9_4", "Fatiga", "PHQ-A", "10 días", "2 puntos", "Más de la mitad"),
    ("phq9_6", "Culpa", "PHQ-A", "6 días", "1 punto", "Algunos días"),
    ("gad7_1", "Ansiedad", "GAD-7", "10 días", "2 puntos", "Más de la mitad"),
    ("gad7_2", "Preocupación incontrolable", "GAD-7", "8 días", "2 puntos", "Más de la mitad"),
    ("gad7_4", "Tensión", "GAD-7", "9 días", "2 puntos", "Más de la mitad"),
]
tabla(
    doc,
    ["Ítem", "Criterio DSM-5", "Módulo", "Días", "Puntos", "Frase oficial"],
    filas_maria,
)

parrafo(
    doc,
    "Sumando los ítems PHQ-A: 2+2+3+2+1 = 10 puntos / 27. Eso cae en "
    "el rango 10-14 = “Moderada”. La acción asociada en el DSM/PHQ "
    "oficial es “alerta al psicólogo”.",
)
parrafo(
    doc,
    "Sumando los ítems GAD-7: 2+2+2 = 6 puntos / 21. Eso cae en el "
    "rango 5-9 = “Leve”. Acción: recomendaciones de autocuidado.",
)

h2(doc, "11.2 Tabla oficial de severidad PHQ-A")
filas_sev_phq = [
    ("0 a 4", "Mínima", "Solo monitoreo"),
    ("5 a 9", "Leve", "Recomendaciones de autocuidado"),
    ("10 a 14", "Moderada", "Alerta al psicólogo"),
    ("15 a 19", "Moderada-severa", "Alerta urgente"),
    ("20 a 27", "Severa", "Protocolo de emergencia"),
]
tabla(doc, ["Puntaje", "Nivel", "Acción sugerida"], filas_sev_phq)
parrafo(
    doc,
    "Fuente: Kroenke & Spitzer (2001), validado en miles de estudios.",
    italic=True,
)

h2(doc, "11.3 Tabla oficial de severidad GAD-7")
filas_sev_gad = [
    ("0 a 4", "Mínima", "Solo monitoreo"),
    ("5 a 9", "Leve", "Recomendaciones de autocuidado"),
    ("10 a 14", "Moderada", "Alerta al psicólogo"),
    ("15 a 21", "Severa", "Alerta urgente"),
]
tabla(doc, ["Puntaje", "Nivel", "Acción sugerida"], filas_sev_gad)
parrafo(
    doc,
    "Fuente: Spitzer & Löwe (2006), validado en miles de estudios.",
    italic=True,
)

h2(doc, "11.4 Confiabilidad del reporte por cobertura")
parrafo(
    doc,
    "Si el alumno escribió pocos días, el reporte se marca con "
    "confiabilidad baja, para que la psicóloga sepa que no debe "
    "tomarlo como base de decisión clínica.",
)
filas_conf = [
    ("Menos de 5 días", "Baja", "Solo orientativo"),
    ("5 a 9 días", "Media", "Útil con cautela"),
    ("10 o más días", "Alta", "Base válida para decisión clínica"),
]
tabla(doc, ["Días escritos en el ciclo", "Confiabilidad", "Uso clínico"], filas_conf)


salto(doc)


# ════════════════════════════════════════════════════════════════════
# 12. POSIBLE RIESGO SEGÚN DSM-5
# ════════════════════════════════════════════════════════════════════
h1(doc, "12. Cómo Sami marca un “posible riesgo según DSM-5”")

parrafo(
    doc,
    "Esta es la parte donde Sami se asoma al manual oficial sin "
    "diagnosticar. La regla operativa es la siguiente: un ítem se "
    "considera “cumplido” cuando obtuvo 2 o más puntos en el ciclo "
    "(es decir, el síntoma estuvo presente más de la mitad de los "
    "días). Esto refleja el criterio clínico estándar: para contar como "
    "síntoma significativo, debe estar la mayor parte del tiempo.",
)

h2(doc, "12.1 Posible riesgo de Episodio Depresivo Mayor (EDM)")
parrafo(doc, "El DSM-5 dice que para hablar de un EDM se necesitan:")
bullet(doc, "Al menos 5 de los 9 síntomas durante 2 semanas.")
bullet(
    doc,
    "Y al menos uno de esos síntomas debe ser ánimo deprimido (ítem 2 "
    "del PHQ-A) o pérdida de interés/anhedonia (ítem 1 del PHQ-A).",
)
parrafo(
    doc,
    "Sami aplica esa misma regla:",
    negrita=True,
)
bullet(
    doc,
    "Si cuenta ≥ 5 ítems del PHQ-A con puntaje ≥ 2 Y al menos uno es "
    "el ítem 1 o el ítem 2 → marca “Posible riesgo de Episodio "
    "Depresivo Mayor (DSM-5)”.",
)
bullet(
    doc,
    "Si no cumple → no marca nada (o marca “Sin indicador”).",
)

h2(doc, "12.2 Posible riesgo de Trastorno de Ansiedad Generalizada (TAG)")
parrafo(doc, "El DSM-5 dice que para hablar de un TAG se necesitan:")
bullet(doc, "Ansiedad excesiva.")
bullet(doc, "Dificultad para controlar la preocupación.")
bullet(doc, "Al menos 3 síntomas asociados durante el período.")
parrafo(
    doc,
    "Sami aplica esa misma regla:",
    negrita=True,
)
bullet(
    doc,
    "Si el ítem 1 del GAD-7 (ansiedad) y el ítem 2 (no logra controlar "
    "la preocupación) están cumplidos, Y en total hay ≥ 3 ítems del "
    "GAD-7 con puntaje ≥ 2 → marca “Posible riesgo de Trastorno de "
    "Ansiedad Generalizada (DSM-5)”.",
)
bullet(doc, "Si no cumple → “Sin indicador”.")

h2(doc, "12.3 Lenguaje cuidadoso, siempre")
parrafo(
    doc,
    "El sistema en ningún lugar dice “el alumno tiene depresión” o "
    "“el alumno tiene ansiedad”. Lo más fuerte que dice es:",
    negrita=True,
)
cita(
    doc,
    "“Posible riesgo de Episodio Depresivo Mayor (DSM-5). 6 de 9 "
    "ítems cumplidos; incluye ánimo deprimido y anhedonia.”",
)
parrafo(
    doc,
    "Y siempre va acompañado de una nota fija que dice:",
)
cita(
    doc,
    "“Esta sección NO es un diagnóstico. Es un indicador de que el "
    "patrón observado en el diario coincide con los criterios mínimos "
    "del DSM-5. El diagnóstico es competencia exclusiva del profesional "
    "de salud mental.”",
)


# ════════════════════════════════════════════════════════════════════
# 13. PROTOCOLO DE CRISIS
# ════════════════════════════════════════════════════════════════════
h1(doc, "13. Qué pasa cuando hay una crisis")

parrafo(
    doc,
    "Sami trata a la crisis como excepción a todas las reglas. Una "
    "crisis se detecta de dos maneras:",
)
bullet(
    doc,
    "El ítem 9 del PHQ-A (pensamientos de hacerse daño) recibe puntaje "
    "≥ 1 en alguna entrada.",
)
bullet(
    doc,
    "La inteligencia artificial detecta un puntaje de “riesgo "
    "suicida” por encima del 40% (umbral más bajo de lo habitual, "
    "para no dejar pasar señales).",
)

h2(doc, "13.1 Lo que el sistema hace en el instante")
numerado(
    doc,
    "Crea un evento de SOS (Sosevent) marcado como “abierto” en el "
    "panel de la psicóloga.",
)
numerado(doc, "Le muestra al alumno en pantalla, sin susto pero con claridad:")
cita(
    doc,
    "“Lo que escribiste me importa. Si lo necesitas ahora, llama a la "
    "Línea 113 (opción 5) o a SALUDLINE 106. No estás solo/a.”",
)
numerado(
    doc,
    "Notifica a la psicóloga inmediatamente (alerta en su dashboard).",
)
numerado(
    doc,
    "Marca el nivel de riesgo del alumno como CRÍTICO, sin importar "
    "el resto del ciclo.",
)

h2(doc, "13.2 Lo que hace la psicóloga")
bullet(doc, "Contacta al alumno apenas ve el SosEvent.")
bullet(
    doc,
    "Activa el protocolo del colegio: contacto con los padres, "
    "evaluación de la necesidad de derivación a salud mental, "
    "documentación del incidente.",
)
bullet(
    doc,
    "Agenda una cita de atención de crisis. Al completarla, el ciclo "
    "en curso del alumno se cierra anticipadamente.",
)


# ════════════════════════════════════════════════════════════════════
# 14. QUÉ VE CADA ROL EN PANTALLA
# ════════════════════════════════════════════════════════════════════
h1(doc, "14. Qué ve cada rol en pantalla")

h2(doc, "14.1 El alumno")
bullet(doc, "Una pantalla de bienvenida con saludo personal de Sami.")
bullet(
    doc,
    "Un cuadro de texto grande para escribir lo que sienta hoy, "
    "con un prompt suave de Sami como guía.",
)
bullet(
    doc,
    "Una pantalla “Mi proceso” donde ve cuántas entradas ha escrito y "
    "qué día va de su ciclo (sin puntajes ni etiquetas).",
)
bullet(
    doc,
    "Un menú de Recursos con líneas de ayuda y consejos de autocuidado.",
)
bullet(doc, "El consentimiento informado y su perfil personal.")
bullet(
    doc,
    "NUNCA: puntajes PHQ-A / GAD-7, etiquetas tipo “tienes depresión”, "
    "ni el detalle clínico del análisis.",
)

h2(doc, "14.2 La psicóloga")

h3(doc, "Panel de monitoreo (dashboard)")
bullet(
    doc,
    "5 tarjetas con números: total de estudiantes, en riesgo CRÍTICO, "
    "ALTO, MEDIO, BAJO/sin evaluar.",
)
bullet(doc, "Lista de alertas tempranas con estudiantes que necesitan revisión.")
bullet(doc, "Lista de próximas citas, con badge de “Crisis” si aplica.")
bullet(
    doc,
    "Tabla de todos los estudiantes con, por cada uno: nombre, "
    "entradas escritas, PHQ-A actual con severidad, GAD-7 actual con "
    "severidad, posibles riesgos DSM-5 (en chips ámbar tipo “Posible "
    "riesgo · Episodio Depresivo Mayor”), nivel de alerta y botones de "
    "acción.",
)

h3(doc, "Historial detallado de un estudiante")
bullet(doc, "Datos del estudiante y nivel actual.")
bullet(
    doc,
    "Gráfico de evolución temporal del nivel de riesgo (línea con "
    "fechas).",
)
bullet(
    doc,
    "Reporte del ciclo actual: tabla con cada ítem (código, módulo "
    "PHQ-A o GAD-7, criterio DSM-5, días con síntoma, puntos, frase "
    "oficial).",
)
bullet(
    doc,
    "Sección “Posibles riesgos según DSM-5”: dos tarjetas (EDM y TAG) "
    "con los ítems cumplidos y el badge “Posible riesgo” o “Sin "
    "indicador”, más el disclaimer fijo.",
)
bullet(
    doc,
    "Lista de entradas del diario, una por una. Por cada entrada se "
    "muestra: fecha, PHQ-A y GAD-7 de esa entrada, ítems activados con "
    "criterio DSM-5 y palabras clave que los activaron. NO se muestra "
    "el texto crudo del alumno (privacidad).",
)
bullet(doc, "Botones para agendar cita, dejar nota clínica privada, enviar mensaje.")

h2(doc, "14.3 El administrador")
bullet(doc, "Gestión de usuarios: alta, baja, activación.")
bullet(doc, "Configuración de preguntas y frecuencia del diario.")
bullet(doc, "Lista de respaldos automáticos de la base de datos.")
bullet(
    doc,
    "Tabla de auditoría: cada llamada al sistema queda registrada con "
    "usuario, rol, endpoint, fecha y resultado.",
)
bullet(doc, "Panel del modelo BERT: estado, dispositivo, botón de recarga.")
bullet(doc, "Sliders para calibrar los umbrales del modelo por condición.")
bullet(doc, "NO ve el contenido de ningún diario.")


# ════════════════════════════════════════════════════════════════════
# 15. PROTECCIÓN DE DATOS Y CONFIDENCIALIDAD
# ════════════════════════════════════════════════════════════════════
h1(doc, "15. Protección de datos y confidencialidad")

bullet(
    doc,
    "Las contraseñas se guardan con hash (no se pueden leer en texto "
    "plano, ni siquiera el administrador puede verlas).",
)
bullet(
    doc,
    "Las sesiones se manejan con JWT (tokens que expiran a las 24 "
    "horas).",
)
bullet(
    doc,
    "El contenido del diario nunca se exporta ni se muestra a los "
    "padres, profesores o autoridades del colegio. Solo a la psicóloga, "
    "y aún ella ve solo el análisis (no el texto crudo).",
)
bullet(
    doc,
    "Cada acceso al sistema queda en una bitácora de auditoría.",
)
bullet(
    doc,
    "El sistema cumple con los principios de la Ley N° 29733 sobre "
    "protección de datos personales en el Perú: minimización de datos, "
    "finalidad específica, consentimiento informado, derecho de acceso "
    "y de rectificación, derecho al olvido.",
)
bullet(
    doc,
    "Se realizan respaldos automáticos diarios de la base de datos.",
)


# ════════════════════════════════════════════════════════════════════
# 16. GLOSARIO
# ════════════════════════════════════════════════════════════════════
h1(doc, "16. Glosario rápido")

filas_glosario = [
    ("Sami", "El nombre amigable del sistema. Lo que el alumno ve."),
    (
        "Diario",
        "El cuaderno digital donde el alumno escribe libremente cómo "
        "se siente. Eje del sistema.",
    ),
    (
        "Ciclo",
        "Período de 14 días desde la primera entrada. Al cierre se "
        "genera el reporte clínico oficial.",
    ),
    (
        "Entrada",
        "Cada texto que el alumno escribe en el diario, asociado a "
        "una fecha.",
    ),
    (
        "Ítem",
        "Cada uno de los 9 síntomas del PHQ-A o los 7 del GAD-7. "
        "Tiene un código (phq9_1, gad7_3…), una palabra clave y un "
        "criterio DSM-5.",
    ),
    (
        "PHQ-A",
        "Patient Health Questionnaire para adolescentes. Cuestionario "
        "de 9 preguntas para medir depresión. Total 0-27.",
    ),
    (
        "GAD-7",
        "Generalized Anxiety Disorder. Cuestionario de 7 preguntas "
        "para medir ansiedad. Total 0-21.",
    ),
    (
        "DSM-5",
        "Manual diagnóstico oficial de los trastornos de salud "
        "mental. Sami solo lo usa para etiquetar criterios y verificar "
        "patrones, no para diagnosticar.",
    ),
    (
        "Likert",
        "Escala de respuesta con 4 niveles fijos: 0 = Nunca, 1 = "
        "Algunos días, 2 = Más de la mitad de los días, 3 = Casi "
        "todos los días.",
    ),
    (
        "BERT / BETO",
        "El modelo de inteligencia artificial que lee el texto del "
        "alumno en español. BETO es la versión en español de BERT "
        "(de Google).",
    ),
    (
        "Zero-shot",
        "Técnica que permite usar el modelo para clasificar sin "
        "tener que entrenarlo de nuevo en cada categoría.",
    ),
    (
        "Crisis",
        "Situación de riesgo grave (ej. ideas de hacerse daño). "
        "Activa protocolo inmediato y alerta a la psicóloga sin "
        "esperar al cierre del ciclo.",
    ),
    (
        "JWT",
        "Tipo de token de sesión que mantiene segura la "
        "autenticación del usuario.",
    ),
    (
        "Cobertura",
        "Porcentaje de días con entrada respecto al total del ciclo. "
        "Determina la confiabilidad del reporte.",
    ),
]
tabla(doc, ["Término", "Qué significa"], filas_glosario)


# ════════════════════════════════════════════════════════════════════
# 17. CIERRE
# ════════════════════════════════════════════════════════════════════
h1(doc, "17. En resumen")

parrafo(
    doc,
    "Sami es un diario digital con un “lector inteligente” detrás que "
    "ayuda a la psicóloga del colegio a darse cuenta a tiempo cuando un "
    "alumno necesita atención. No reemplaza a la psicóloga, no "
    "diagnostica, no etiqueta. Solo escucha al alumno con respeto, "
    "cuenta las señales que aparecen en lo que escribe, y prepara un "
    "reporte ordenado para que la profesional decida qué hacer.",
)

parrafo(
    doc,
    "Las decisiones clave del sistema están respaldadas por instrumentos "
    "clínicos validados (PHQ-A, GAD-7), el manual oficial DSM-5, y "
    "buenas prácticas de protección de datos personales (Ley 29733). "
    "Cada vez que el sistema dice algo, lo dice con palabras del manual, "
    "no inventadas.",
)

parrafo(
    doc,
    "El alumno tiene un espacio seguro, los padres tienen un sistema "
    "transparente, la psicóloga tiene una herramienta de monitoreo, y "
    "el colegio cumple con su deber de cuidado.",
    italic=True,
)


# ════════════════════════════════════════════════════════════════════
# GUARDAR
# ════════════════════════════════════════════════════════════════════
import os
OUT = os.path.join(os.path.dirname(__file__), "Sami_para_asesor.docx")
doc.save(OUT)
print(f"✅ Documento generado: {OUT}")
