"""
Documento de fuentes científicas y normativas del sistema Sami.
Salida: docs/Sami_fuentes.docx

Cada dato del documento del asesor está respaldado por una fuente
académica revisada por pares o por normativa oficial peruana. Este
archivo agrupa todas las fuentes por temática, en formato APA 7, e
indica en qué parte del sistema o del documento del asesor se usa.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

VERDE = RGBColor(0x05, 0x96, 0x69)
VERDE_OSC = RGBColor(0x04, 0x6A, 0x49)
GRIS_OSCURO = RGBColor(0x0A, 0x0A, 0x0A)
GRIS_MEDIO = RGBColor(0x5A, 0x5A, 0x5A)


def h1(doc, texto):
    p = doc.add_heading(texto, level=1)
    for run in p.runs:
        run.font.color.rgb = VERDE
        run.font.size = Pt(20)


def h2(doc, texto):
    p = doc.add_heading(texto, level=2)
    for run in p.runs:
        run.font.color.rgb = VERDE_OSC
        run.font.size = Pt(15)


def h3(doc, texto):
    p = doc.add_heading(texto, level=3)
    for run in p.runs:
        run.font.color.rgb = GRIS_OSCURO
        run.font.size = Pt(12)


def parrafo(doc, texto, negrita=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(size)
    if negrita:
        run.bold = True
    if italic:
        run.italic = True
    return p


def ref_apa(doc, texto):
    """Referencia APA con sangría francesa."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(texto)
    run.font.size = Pt(11)
    return p


def uso(doc, texto):
    """Línea que indica dónde se usa la fuente."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Uso en el sistema: ")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = VERDE_OSC
    run2 = p.add_run(texto)
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRIS_MEDIO


def bullet(doc, texto, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(texto)
    run.font.size = Pt(size)


def salto(doc):
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


# ── PORTADA ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\nSami")
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = VERDE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fuentes científicas y normativas")
run.italic = True
run.font.size = Pt(16)
run.font.color.rgb = GRIS_MEDIO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "\n\nRespaldo bibliográfico de cada dato, escala, regla y "
    "decisión clínica del sistema.\nFormato APA 7.\n\n"
    "Universidad Peruana de Ciencias Aplicadas — UPC"
)
run.font.size = Pt(11)
run.font.color.rgb = GRIS_OSCURO

salto(doc)


# ════════════════════════════════════════════════════════════════════
# INTRODUCCIÓN
# ════════════════════════════════════════════════════════════════════
h1(doc, "Introducción")
parrafo(
    doc,
    "Este documento reúne todas las fuentes que sostienen cada "
    "afirmación clínica, técnica y normativa del sistema Sami. No se "
    "ha utilizado contenido enciclopédico de divulgación abierta "
    "(Wikipedia, blogs no revisados) como respaldo de ninguna "
    "decisión. Las fuentes citadas son: (a) artículos en revistas "
    "indexadas con revisión por pares; (b) manuales clínicos "
    "oficiales (DSM-5 de la APA); (c) normativa peruana publicada en "
    "el Diario Oficial El Peruano; y (d) recursos institucionales del "
    "MINSA y MINEDU.",
)
parrafo(
    doc,
    "Cada fuente está acompañada de una línea que indica exactamente "
    "qué parte del sistema o del documento de defensa la utiliza, "
    "para que el lector pueda trazar la procedencia de cada dato.",
)


# ════════════════════════════════════════════════════════════════════
# 1. PHQ-A / PHQ-9 (DEPRESIÓN)
# ════════════════════════════════════════════════════════════════════
h1(doc, "1. PHQ-A y PHQ-9 — Escala de depresión")

h3(doc, "1.1 Validación original del PHQ-9")
ref_apa(
    doc,
    "Kroenke, K., Spitzer, R. L., & Williams, J. B. W. (2001). The "
    "PHQ-9: Validity of a brief depression severity measure. "
    "Journal of General Internal Medicine, 16(9), 606–613. "
    "https://doi.org/10.1046/j.1525-1497.2001.016009606.x",
)
uso(
    doc,
    "Cuestionario base de 9 ítems para depresión. Es la fuente "
    "primaria del PHQ-9. Define los 9 ítems, las 4 opciones de "
    "respuesta (0–3) y el rango total 0–27.",
)

h3(doc, "1.2 Puntos de corte de severidad del PHQ-9")
ref_apa(
    doc,
    "Kroenke, K., & Spitzer, R. L. (2002). The PHQ-9: A new "
    "depression diagnostic and severity measure. Psychiatric "
    "Annals, 32(9), 509–515. "
    "https://doi.org/10.3928/0048-5713-20020901-06",
)
uso(
    doc,
    "Tabla de severidad usada en `app/config.py` y mostrada al "
    "psicólogo en `StudentHistoryView`: Mínima (0–4), Leve (5–9), "
    "Moderada (10–14), Moderada-severa (15–19), Severa (20–27).",
)

h3(doc, "1.3 PHQ-A — Validación en adolescentes")
ref_apa(
    doc,
    "Johnson, J. G., Harris, E. S., Spitzer, R. L., & Williams, "
    "J. B. W. (2002). The Patient Health Questionnaire for "
    "Adolescents: Validation of an instrument for the assessment of "
    "mental disorders among adolescent primary care patients. "
    "Journal of Adolescent Health, 30(3), 196–204. "
    "https://doi.org/10.1016/S1054-139X(01)00333-0",
)
uso(
    doc,
    "Fuente clave del proyecto. De aquí provienen las 4 frases "
    "oficiales (Nunca / Algunos días / Más de la mitad de los días / "
    "Casi todos los días) cuya traducción literal es la base de la "
    "tabla días→puntos del ciclo de 14 días. Implementada en "
    "`app/services/diario_analisis_service.py` función "
    "`_dias_a_puntos`. También en el documento del asesor "
    "Sección 9 y Apéndice de tablas.",
)

h3(doc, "1.4 Validación peruana del PHQ-9")
ref_apa(
    doc,
    "Calderón, M., Gálvez-Buccollini, J. A., Cueva, G., Ordoñez, "
    "C., Bromley, C., & Fiestas, F. (2012). Validación de la "
    "versión peruana del PHQ-9 para el diagnóstico de depresión. "
    "Revista Peruana de Medicina Experimental y Salud Pública, "
    "29(4), 578–579.",
)
uso(
    doc,
    "Respalda el uso del PHQ-9 en población peruana. Necesario para "
    "defender la aplicabilidad de la escala en el contexto local "
    "del colegio.",
)

ref_apa(
    doc,
    "Saravia, J. C., Iberico, C., & Yearwood, K. (2014). "
    "Validación del cuestionario de salud del paciente-9 (PHQ-9) "
    "para detectar trastornos depresivos en estudiantes "
    "universitarios peruanos. Revista de Psicología (PUCP), "
    "32(2), 254–272.",
)
uso(
    doc,
    "Validación específica en estudiantes peruanos. Aporte directo "
    "para defender la aplicación en una población escolar/joven en "
    "Perú.",
)


# ════════════════════════════════════════════════════════════════════
# 2. GAD-7 (ANSIEDAD)
# ════════════════════════════════════════════════════════════════════
h1(doc, "2. GAD-7 — Escala de ansiedad")

h3(doc, "2.1 Validación original del GAD-7")
ref_apa(
    doc,
    "Spitzer, R. L., Kroenke, K., Williams, J. B. W., & Löwe, B. "
    "(2006). A brief measure for assessing generalized anxiety "
    "disorder: The GAD-7. Archives of Internal Medicine, 166(10), "
    "1092–1097. https://doi.org/10.1001/archinte.166.10.1092",
)
uso(
    doc,
    "Fuente primaria del GAD-7. Define los 7 ítems, las 4 opciones "
    "Likert y los puntos de corte de severidad (Mínima 0–4, Leve "
    "5–9, Moderada 10–14, Severa 15–21) que el sistema aplica al "
    "cierre de cada ciclo. Implementado en "
    "`app/config.py: GAD7_SEVERIDAD`.",
)

h3(doc, "2.2 Estandarización en población general")
ref_apa(
    doc,
    "Löwe, B., Decker, O., Müller, S., Brähler, E., Schellberg, "
    "D., Herzog, W., & Herzberg, P. Y. (2008). Validation and "
    "standardization of the Generalized Anxiety Disorder Screener "
    "(GAD-7) in the general population. Medical Care, 46(3), "
    "266–274. https://doi.org/10.1097/MLR.0b013e318160d093",
)
uso(
    doc,
    "Respalda el uso del GAD-7 fuera de atención primaria, "
    "incluyendo poblaciones generales y estudios comunitarios.",
)

h3(doc, "2.3 Versión en español")
ref_apa(
    doc,
    "García-Campayo, J., Zamorano, E., Ruiz, M. A., Pardo, A., "
    "Pérez-Páramo, M., López-Gómez, V., Freire, O., & Rejas, J. "
    "(2010). Cultural adaptation into Spanish of the generalized "
    "anxiety disorder-7 (GAD-7) scale as a screening tool. Health "
    "and Quality of Life Outcomes, 8, 8. "
    "https://doi.org/10.1186/1477-7525-8-8",
)
uso(
    doc,
    "Versión en español culturalmente adaptada del GAD-7. Base de "
    "la formulación de los 7 ítems en español usada en Sami "
    "(ver `PHQ9_ITEMS` y `GAD7_ITEMS` en `app/config.py`).",
)


# ════════════════════════════════════════════════════════════════════
# 3. DSM-5
# ════════════════════════════════════════════════════════════════════
h1(doc, "3. DSM-5 — Manual diagnóstico")

h3(doc, "3.1 Manual oficial")
ref_apa(
    doc,
    "American Psychiatric Association. (2013). Diagnostic and "
    "Statistical Manual of Mental Disorders (5th ed.). American "
    "Psychiatric Publishing. "
    "https://doi.org/10.1176/appi.books.9780890425596",
)
uso(
    doc,
    "Fuente oficial de TODOS los criterios diagnósticos citados en "
    "el sistema. Los 9 ítems del PHQ-A están etiquetados con su "
    "criterio DSM-5 (campo `criterio_dsm5` en `app/config.py`). "
    "Las reglas de 'posible riesgo de Episodio Depresivo Mayor' y "
    "'posible riesgo de Trastorno de Ansiedad Generalizada' usan "
    "los umbrales mínimos definidos aquí.",
)

h3(doc, "3.2 Edición revisada (DSM-5-TR)")
ref_apa(
    doc,
    "American Psychiatric Association. (2022). Diagnostic and "
    "Statistical Manual of Mental Disorders, Text Revision "
    "(DSM-5-TR). American Psychiatric Publishing. "
    "https://doi.org/10.1176/appi.books.9780890425787",
)
uso(
    doc,
    "Edición vigente que mantiene los criterios usados en Sami "
    "para Episodio Depresivo Mayor (criterios A1–A9, requiere "
    "≥5 síntomas durante 2 semanas con al menos ánimo deprimido o "
    "anhedonia) y Trastorno de Ansiedad Generalizada (ansiedad y "
    "preocupación excesivas con dificultad de control + ≥3 "
    "síntomas asociados).",
)

h3(doc, "3.3 Criterios específicos de Episodio Depresivo Mayor")
parrafo(
    doc,
    "Páginas 160-161 del DSM-5 (Sección II — Trastornos depresivos):",
    italic=True,
)
bullet(
    doc,
    "A1: Estado de ánimo deprimido la mayor parte del día, casi "
    "todos los días.",
)
bullet(doc, "A2: Pérdida de interés o placer (anhedonia).")
bullet(doc, "A3: Cambio significativo de peso o apetito.")
bullet(doc, "A4: Insomnio o hipersomnia.")
bullet(doc, "A5: Agitación o enlentecimiento psicomotor observables.")
bullet(doc, "A6: Fatiga o pérdida de energía.")
bullet(doc, "A7: Sentimientos de inutilidad o culpa excesiva.")
bullet(doc, "A8: Dificultad para concentrarse o indecisión.")
bullet(doc, "A9: Pensamientos recurrentes de muerte o ideación suicida.")
parrafo(
    doc,
    "Umbral mínimo para sospecha clínica: ≥5 síntomas durante un "
    "período de 2 semanas, al menos uno debe ser A1 o A2.",
)
uso(
    doc,
    "Estos 9 criterios están mapeados 1:1 a los 9 ítems del PHQ-A "
    "en el sistema. La función `_evaluar_criterios_dsm5` en "
    "`app/services/diario_analisis_service.py` aplica exactamente "
    "este umbral.",
)

h3(doc, "3.4 Criterios específicos de Trastorno de Ansiedad Generalizada")
parrafo(doc, "Páginas 222 del DSM-5 (Sección II — Trastornos de ansiedad):", italic=True)
bullet(
    doc,
    "A: Ansiedad y preocupación excesivas la mayor parte de los "
    "días durante al menos 6 meses.",
)
bullet(doc, "B: La persona tiene dificultad para controlar la preocupación.")
bullet(
    doc,
    "C: Al menos 3 síntomas asociados de los 6 listados: "
    "inquietud, fatigabilidad, dificultad para concentrarse, "
    "irritabilidad, tensión muscular, alteraciones del sueño.",
)
uso(
    doc,
    "Implementado en la regla de TAG del sistema: requiere "
    "GAD-7 ítem 1 (ansiedad) + ítem 2 (descontrol de "
    "preocupación) cumplidos, más ≥3 ítems totales del GAD-7 con "
    "puntaje ≥2.",
)


# ════════════════════════════════════════════════════════════════════
# 4. ESCALAS COMPLEMENTARIAS
# ════════════════════════════════════════════════════════════════════
h1(doc, "4. Escalas complementarias")

h3(doc, "4.1 C-SSRS — Columbia Suicide Severity Rating Scale")
ref_apa(
    doc,
    "Posner, K., Brown, G. K., Stanley, B., Brent, D. A., Yershova, "
    "K. V., Oquendo, M. A., Currier, G. W., Melvin, G. A., Greenhill, "
    "L., Shen, S., & Mann, J. J. (2011). The Columbia–Suicide "
    "Severity Rating Scale: Initial validity and internal "
    "consistency findings from three multisite studies with "
    "adolescents and adults. American Journal of Psychiatry, "
    "168(12), 1266–1277. "
    "https://doi.org/10.1176/appi.ajp.2011.10111704",
)
uso(
    doc,
    "Referencia del estándar internacional para evaluación de "
    "riesgo suicida. El umbral más sensible (40%) del modelo BERT "
    "para 'riesgo_suicida' en `app/config.py` se justifica con "
    "esta escala: en clínica de adolescentes se prefiere alta "
    "sensibilidad sobre alta especificidad. También respalda el "
    "criterio del PHQ-A ítem 9 que dispara protocolo de crisis "
    "inmediato.",
)

h3(doc, "4.2 ASRS v1.1 — Adult ADHD Self-Report Scale (OMS)")
ref_apa(
    doc,
    "Kessler, R. C., Adler, L., Ames, M., Demler, O., Faraone, S., "
    "Hiripi, E., Howes, M. J., Jin, R., Secnik, K., Spencer, T., "
    "Ustun, T. B., & Walters, E. E. (2005). The World Health "
    "Organization Adult ADHD Self-Report Scale (ASRS): A short "
    "screening scale for use in the general population. "
    "Psychological Medicine, 35(2), 245–256. "
    "https://doi.org/10.1017/S0033291704002892",
)
uso(
    doc,
    "Base teórica de la categoría 'tdah' en `CONDICIONES` del "
    "modelo BERT. Aunque la categoría está disponible en el "
    "monitoreo, las reglas DSM-5 implementadas se limitan a EDM "
    "y TAG.",
)

h3(doc, "4.3 UCLA-3 — Escala breve de soledad")
ref_apa(
    doc,
    "Hughes, M. E., Waite, L. J., Hawkley, L. C., & Cacioppo, J. T. "
    "(2004). A short scale for measuring loneliness in large "
    "surveys: Results from two population-based studies. Research "
    "on Aging, 26(6), 655–672. "
    "https://doi.org/10.1177/0164027504268574",
)
uso(
    doc,
    "Base de la categoría 'soledad' del sistema. Aplicable al "
    "contexto adolescente, validado internacionalmente.",
)


# ════════════════════════════════════════════════════════════════════
# 5. MODELO BERT / BETO / XNLI
# ════════════════════════════════════════════════════════════════════
h1(doc, "5. Modelo de procesamiento de lenguaje natural")

h3(doc, "5.1 BERT — Modelo base")
ref_apa(
    doc,
    "Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). "
    "BERT: Pre-training of deep bidirectional transformers for "
    "language understanding. In Proceedings of the 2019 Conference "
    "of the North American Chapter of the Association for "
    "Computational Linguistics: Human Language Technologies "
    "(NAACL-HLT), Vol. 1, pp. 4171–4186. "
    "https://doi.org/10.18653/v1/N19-1423",
)
uso(
    doc,
    "Arquitectura base del modelo usado en Sami. Publicación "
    "original de Google AI Research.",
)

h3(doc, "5.2 BETO — BERT en español")
ref_apa(
    doc,
    "Cañete, J., Chaperon, G., Fuentes, R., Ho, J.-H., Kang, H., & "
    "Pérez, J. (2020). Spanish pre-trained BERT model and "
    "evaluation data. In PML4DC at ICLR 2020 (Practical ML for "
    "Developing Countries Workshop). "
    "https://arxiv.org/abs/2308.02976",
)
uso(
    doc,
    "Fuente del modelo BERT entrenado en español. Universidad de "
    "Chile. Es la base del modelo que Sami usa.",
)

h3(doc, "5.3 XNLI — Inferencia interlingüística")
ref_apa(
    doc,
    "Conneau, A., Lample, G., Rinott, R., Williams, A., Bowman, S. "
    "R., Schwenk, H., & Stoyanov, V. (2018). XNLI: Evaluating "
    "cross-lingual sentence representations. In Proceedings of the "
    "2018 Conference on Empirical Methods in Natural Language "
    "Processing (EMNLP), pp. 2475–2485. "
    "https://doi.org/10.18653/v1/D18-1269",
)
uso(
    doc,
    "Base teórica del fine-tuning XNLI que permite clasificación "
    "zero-shot en español.",
)

h3(doc, "5.4 Zero-shot text classification por entailment")
ref_apa(
    doc,
    "Yin, W., Hay, J., & Roth, D. (2019). Benchmarking zero-shot "
    "text classification: Datasets, evaluation and entailment "
    "approach. In Proceedings of the 2019 Conference on Empirical "
    "Methods in Natural Language Processing and the 9th "
    "International Joint Conference on Natural Language Processing "
    "(EMNLP-IJCNLP), pp. 3914–3923. "
    "https://doi.org/10.18653/v1/D19-1404",
)
uso(
    doc,
    "Sostiene metodológicamente el enfoque de hypothesis templating "
    "(\"Este texto expresa {}.\") implementado en "
    "`app/services/nlp_service.py`.",
)

h3(doc, "5.5 Modelo específico utilizado")
ref_apa(
    doc,
    "Recognai. (2020). bert-base-spanish-wwm-cased-xnli "
    "[Model card]. Hugging Face. "
    "https://huggingface.co/Recognai/"
    "bert-base-spanish-wwm-cased-xnli",
)
uso(
    doc,
    "El identificador exacto del modelo descargado por Sami "
    "(`MODEL_NAME` en `app/config.py`). Combina BETO (Cañete et "
    "al., 2020) con fine-tuning en XNLI español.",
)


# ════════════════════════════════════════════════════════════════════
# 6. NORMATIVA PERUANA Y CONSENTIMIENTO
# ════════════════════════════════════════════════════════════════════
h1(doc, "6. Normativa peruana y consentimiento informado")

h3(doc, "6.1 Ley de Protección de Datos Personales")
ref_apa(
    doc,
    "Congreso de la República del Perú. (2011, 3 de julio). "
    "Ley N° 29733, Ley de Protección de Datos Personales. Diario "
    "Oficial El Peruano. "
    "https://www.minjus.gob.pe/wp-content/uploads/2013/04/"
    "LEY-29733.pdf",
)
uso(
    doc,
    "Marco legal del manejo de datos personales del alumno. "
    "Citado en `ConsentView.vue` y en el documento del asesor "
    "Sección 5 (consentimiento) y Sección 15 (protección de datos).",
)

h3(doc, "6.2 Reglamento de la Ley 29733")
ref_apa(
    doc,
    "Ministerio de Justicia y Derechos Humanos del Perú. (2013, "
    "22 de marzo). Decreto Supremo N° 003-2013-JUS, Reglamento de "
    "la Ley N° 29733 de Protección de Datos Personales. Diario "
    "Oficial El Peruano.",
)
uso(
    doc,
    "Detalla principios operativos que el sistema cumple: "
    "consentimiento informado, finalidad específica, "
    "minimización, calidad, seguridad, disposición de datos, "
    "derecho de acceso, rectificación y cancelación.",
)

h3(doc, "6.3 Lineamientos del MINEDU sobre bienestar socioemocional")
ref_apa(
    doc,
    "Ministerio de Educación del Perú. (2020). Resolución "
    "Viceministerial N° 212-2020-MINEDU. Lineamientos para la "
    "Gestión del Bienestar Socioemocional en las Instituciones "
    "Educativas. Lima: MINEDU.",
)
uso(
    doc,
    "Respalda el rol del psicólogo escolar y la necesidad de un "
    "sistema de monitoreo socioemocional. Justifica la pertinencia "
    "del proyecto en el sistema educativo peruano.",
)

ref_apa(
    doc,
    "Ministerio de Educación del Perú. (2021). Cartilla de "
    "Orientación: Acompañamiento socioemocional para estudiantes. "
    "Dirección de Educación Secundaria. Lima: MINEDU.",
)
uso(
    doc,
    "Recursos oficiales para el tipo de acompañamiento al que el "
    "sistema apoya, no reemplaza.",
)

h3(doc, "6.4 Consentimiento informado en adolescentes")
ref_apa(
    doc,
    "Consejo Nacional de Bioética del Perú. (2019). Lineamientos "
    "éticos en investigación con adolescentes. Instituto Nacional "
    "de Salud — INS.",
)
uso(
    doc,
    "Marco ético del consentimiento parental + asentimiento del "
    "menor que el sistema implementa en `ConsentView.vue`.",
)


# ════════════════════════════════════════════════════════════════════
# 7. LÍNEAS DE AYUDA OFICIALES
# ════════════════════════════════════════════════════════════════════
h1(doc, "7. Líneas de ayuda oficiales en el Perú")

h3(doc, "7.1 Línea 113 — MINSA, opción 5 Salud Mental")
ref_apa(
    doc,
    "Ministerio de Salud del Perú. (2024). Línea 113 — Servicio "
    "InfoSalud. Opción 5: Salud mental, atención en crisis "
    "emocional. https://www.gob.pe/minsa",
)
uso(
    doc,
    "Línea de emergencia mostrada al alumno cuando el sistema "
    "detecta crisis (ítem 9 PHQ-A ≥1 o BERT riesgo suicida "
    "≥40%). Implementado en `app/services/sos_service.py` y en "
    "`RecursosView.vue`.",
)

h3(doc, "7.2 SALUDLINE 106 — INSM Honorio Delgado–Hideyo Noguchi")
ref_apa(
    doc,
    "Instituto Nacional de Salud Mental Honorio Delgado–Hideyo "
    "Noguchi. (2024). SALUDLINE 106 — Orientación telefónica en "
    "salud mental. Ministerio de Salud del Perú.",
)
uso(
    doc,
    "Segunda línea de apoyo mostrada al alumno y en la pantalla "
    "de consentimiento. Atención especializada en salud mental.",
)

h3(doc, "7.3 Teléfono de la Esperanza")
ref_apa(
    doc,
    "Teléfono de la Esperanza del Perú. (2024). Línea de apoyo "
    "emocional 24 horas. https://telefonodelaesperanza.org/pe",
)
uso(
    doc,
    "Tercera línea de apoyo, ONG con cobertura 24 horas, citada "
    "en `RecursosView.vue`.",
)


# ════════════════════════════════════════════════════════════════════
# 8. SALUD MENTAL ADOLESCENTE — CONTEXTO
# ════════════════════════════════════════════════════════════════════
h1(doc, "8. Contexto de salud mental adolescente")

h3(doc, "8.1 Datos epidemiológicos globales")
ref_apa(
    doc,
    "World Health Organization. (2021). Adolescent mental health "
    "fact sheet. Geneva: WHO. "
    "https://www.who.int/news-room/fact-sheets/detail/"
    "adolescent-mental-health",
)
uso(
    doc,
    "Justifica la urgencia del problema: 1 de cada 7 adolescentes "
    "(10–19 años) padece un trastorno mental. La depresión es una "
    "de las principales causas de enfermedad y discapacidad.",
)

h3(doc, "8.2 Datos en Perú")
ref_apa(
    doc,
    "Instituto Nacional de Salud Mental Honorio Delgado–Hideyo "
    "Noguchi. (2013). Estudio Epidemiológico de Salud Mental en "
    "Adolescentes en Lima Metropolitana y Callao. Anales de Salud "
    "Mental, 28(Supl. 1).",
)
uso(
    doc,
    "Datos locales que sostienen la pertinencia del proyecto en "
    "el contexto peruano.",
)

ref_apa(
    doc,
    "Defensoría del Pueblo del Perú. (2018). El derecho a la "
    "salud mental: Supervisión de la implementación de la "
    "política pública de atención comunitaria y el respeto a los "
    "derechos humanos. Informe Defensorial N° 180.",
)
uso(
    doc,
    "Marco de derechos humanos asociado al acceso a salud mental "
    "en Perú.",
)


# ════════════════════════════════════════════════════════════════════
# 9. SALUD MENTAL DIGITAL
# ════════════════════════════════════════════════════════════════════
h1(doc, "9. Salud mental digital y apps de bienestar")

h3(doc, "9.1 Eficacia de apps en salud mental")
ref_apa(
    doc,
    "Linardon, J., Cuijpers, P., Carlbring, P., Messer, M., & "
    "Fuller-Tyszkiewicz, M. (2019). The efficacy of app-supported "
    "smartphone interventions for mental health problems: A "
    "meta-analysis of randomized controlled trials. World "
    "Psychiatry, 18(3), 325–336. "
    "https://doi.org/10.1002/wps.20673",
)
uso(
    doc,
    "Meta-análisis que respalda la efectividad de las "
    "intervenciones digitales en salud mental. Defensa contra el "
    "argumento 'una app no sirve'.",
)

h3(doc, "9.2 Psiquiatría digital actual")
ref_apa(
    doc,
    "Torous, J., Bucci, S., Bell, I. H., Kessing, L. V., Faurholt-"
    "Jepsen, M., Whelan, P., Carvalho, A. F., Keshavan, M., "
    "Linardon, J., & Firth, J. (2021). The growing field of "
    "digital psychiatry: Current evidence and the future of apps, "
    "social media, chatbots, and virtual reality. World "
    "Psychiatry, 20(3), 318–335. "
    "https://doi.org/10.1002/wps.20883",
)
uso(
    doc,
    "Estado del arte de la psiquiatría digital. Define las "
    "buenas prácticas que Sami sigue: no diagnostica, complementa "
    "al profesional, respeta privacidad.",
)

h3(doc, "9.3 Diarios emocionales como intervención")
ref_apa(
    doc,
    "Pennebaker, J. W., & Smyth, J. M. (2016). Opening up by "
    "writing it down: How expressive writing improves health and "
    "eases emotional pain (3rd ed.). The Guilford Press.",
)
uso(
    doc,
    "Fundamento teórico de la escritura expresiva como "
    "herramienta terapéutica. Justifica el formato de diario en "
    "Sami por encima de un cuestionario cerrado.",
)


# ════════════════════════════════════════════════════════════════════
# 10. PROCESAMIENTO DE LENGUAJE EN SALUD MENTAL
# ════════════════════════════════════════════════════════════════════
h1(doc, "10. PLN aplicado a salud mental")

h3(doc, "10.1 Detección automática de depresión en texto")
ref_apa(
    doc,
    "Coppersmith, G., Dredze, M., & Harman, C. (2014). Quantifying "
    "mental health signals in Twitter. In Proceedings of the "
    "Workshop on Computational Linguistics and Clinical Psychology: "
    "From Linguistic Signal to Clinical Reality (pp. 51–60). "
    "Association for Computational Linguistics. "
    "https://doi.org/10.3115/v1/W14-3207",
)
uso(
    doc,
    "Demuestra que las señales lingüísticas de salud mental son "
    "detectables computacionalmente. Sostiene el enfoque "
    "metodológico de Sami.",
)

ref_apa(
    doc,
    "Cohan, A., Desmet, B., Yates, A., Soldaini, L., MacAvaney, "
    "S., & Goharian, N. (2018). SMHD: A large-scale resource for "
    "exploring online language usage for multiple mental health "
    "conditions. In Proceedings of the 27th International "
    "Conference on Computational Linguistics (pp. 1485–1497).",
)
uso(
    doc,
    "Dataset y enfoque que respalda el monitoreo multi-condición "
    "(depresión, ansiedad, TDAH, soledad) que Sami implementa.",
)

h3(doc, "10.2 Detección de riesgo suicida en texto")
ref_apa(
    doc,
    "Coppersmith, G., Leary, R., Crutchley, P., & Fine, A. (2018). "
    "Natural language processing of social media as screening for "
    "suicide risk. Biomedical Informatics Insights, 10, "
    "1178222618792860. "
    "https://doi.org/10.1177/1178222618792860",
)
uso(
    doc,
    "Sostiene la viabilidad clínica de detectar riesgo suicida "
    "vía PLN. Respalda el umbral sensible de 40% configurado en "
    "Sami para la categoría 'riesgo_suicida' del modelo BERT.",
)


# ════════════════════════════════════════════════════════════════════
# 11. PROTECCIÓN DE DATOS Y CIBERSEGURIDAD
# ════════════════════════════════════════════════════════════════════
h1(doc, "11. Seguridad de la información")

h3(doc, "11.1 OWASP Top 10")
ref_apa(
    doc,
    "Open Web Application Security Project. (2021). OWASP Top 10: "
    "The ten most critical web application security risks. "
    "https://owasp.org/Top10/",
)
uso(
    doc,
    "Marco de referencia para las decisiones de seguridad de la "
    "arquitectura: hash de contraseñas con bcrypt, autenticación "
    "JWT con expiración de 24 horas, sanitización de inputs, "
    "control de acceso por rol.",
)

h3(doc, "11.2 JWT — Estándar de tokens")
ref_apa(
    doc,
    "Jones, M., Bradley, J., & Sakimura, N. (2015). JSON Web Token "
    "(JWT). Request for Comments: 7519. Internet Engineering Task "
    "Force (IETF). https://datatracker.ietf.org/doc/html/rfc7519",
)
uso(
    doc,
    "Especificación técnica del mecanismo de autenticación "
    "implementado en `app/core/security.py`.",
)


# ════════════════════════════════════════════════════════════════════
# 12. TECNOLOGÍAS USADAS
# ════════════════════════════════════════════════════════════════════
h1(doc, "12. Documentación oficial de las tecnologías")

ref_apa(
    doc,
    "Ramírez, S. (2024). FastAPI documentation. "
    "https://fastapi.tiangolo.com/",
)
uso(doc, "Framework web del backend.")

ref_apa(
    doc,
    "You, E. (2024). Vue.js 3 documentation. "
    "https://vuejs.org/",
)
uso(doc, "Framework frontend.")

ref_apa(
    doc,
    "Hugging Face. (2024). Transformers library documentation. "
    "https://huggingface.co/docs/transformers/",
)
uso(
    doc,
    "Biblioteca usada para cargar y servir el modelo BERT en Sami.",
)

ref_apa(
    doc,
    "Bayer, M. (2024). SQLAlchemy documentation. "
    "https://docs.sqlalchemy.org/",
)
uso(doc, "ORM del backend para SQLite.")


# ════════════════════════════════════════════════════════════════════
# 13. CIERRE
# ════════════════════════════════════════════════════════════════════
salto(doc)
h1(doc, "Notas finales")

parrafo(
    doc,
    "Resumen de respaldo por componente del sistema:",
    negrita=True,
)
bullet(
    doc,
    "Escalas clínicas PHQ-A y GAD-7: validación original "
    "(Kroenke et al., 2001; Spitzer et al., 2006) + validación "
    "en adolescentes (Johnson et al., 2002) + adaptación al "
    "Perú (Calderón et al., 2012; Saravia et al., 2014) + versión "
    "en español (García-Campayo et al., 2010).",
)
bullet(
    doc,
    "Criterios diagnósticos: DSM-5 (APA, 2013) y DSM-5-TR "
    "(APA, 2022) directamente como manuales oficiales.",
)
bullet(
    doc,
    "Tabla días→puntos (Johnson): traducción literal de las "
    "4 frases del cuestionario oficial (Johnson et al., 2002).",
)
bullet(
    doc,
    "Modelo de IA: tres referencias en cascada: BERT (Devlin et "
    "al., 2019) → BETO en español (Cañete et al., 2020) → fine-"
    "tuning XNLI (Conneau et al., 2018; Yin et al., 2019). El "
    "modelo específico es de Recognai (2020).",
)
bullet(
    doc,
    "Riesgo suicida: C-SSRS (Posner et al., 2011) + evidencia de "
    "PLN aplicado (Coppersmith et al., 2018).",
)
bullet(
    doc,
    "Marco legal: Ley 29733 + Reglamento + RVM 212-2020-MINEDU + "
    "Lineamientos del Consejo Nacional de Bioética.",
)
bullet(
    doc,
    "Líneas de ayuda: directamente del MINSA, INSM Honorio "
    "Delgado–Hideyo Noguchi y ONG Teléfono de la Esperanza.",
)
bullet(
    doc,
    "Seguridad: OWASP Top 10 y RFC 7519 (JWT).",
)

parrafo(
    doc,
    "Ninguna afirmación clínica, estadística o normativa del "
    "sistema Sami se basa en fuentes no revisadas por pares o no "
    "oficiales. Las referencias citadas en este documento son "
    "directamente verificables a través de los DOI y enlaces "
    "institucionales proporcionados.",
    italic=True,
)


# ════════════════════════════════════════════════════════════════════
import os
OUT = os.path.join(os.path.dirname(__file__), "Sami_fuentes.docx")
doc.save(OUT)
print(f"✅ Documento generado: {OUT}")
