"""
Documento Word del sistema Sami — versión defensa.
Escrito en lenguaje sencillo para que cualquier persona (psicóloga, asesor de
tesis, padre de familia, jurado) pueda leerlo de corrido y entender el
proyecto completo: qué hace Sami, cómo lo hace, quiénes lo usan, cómo
funcionan los ciclos del diario, qué ve la psicóloga, qué pasa en crisis y
cómo se va a implementar. Apéndices con ítems PHQ-A / GAD-7 y palabras
adaptadas.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

VERDE = RGBColor(0x05, 0x96, 0x69)
GRIS_OSCURO = RGBColor(0x0A, 0x0A, 0x0A)
GRIS_MEDIO = RGBColor(0x5A, 0x5A, 0x5A)
AMBAR = RGBColor(0xB8, 0x85, 0x0A)


def h1(doc, texto):
    p = doc.add_heading(texto, level=1)
    for run in p.runs:
        run.font.color.rgb = VERDE
        run.font.size = Pt(18)


def h2(doc, texto):
    p = doc.add_heading(texto, level=2)
    for run in p.runs:
        run.font.color.rgb = GRIS_OSCURO
        run.font.size = Pt(14)


def h3(doc, texto):
    p = doc.add_heading(texto, level=3)
    for run in p.runs:
        run.font.color.rgb = VERDE
        run.font.size = Pt(12)


def parrafo(doc, texto, negrita=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(size)
    if negrita:
        run.bold = True
    if italic:
        run.italic = True
    return p


def cita(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(texto)
    run.italic = True
    run.font.color.rgb = GRIS_MEDIO


def bullet(doc, texto):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(texto)
    run.font.size = Pt(11)


def alerta(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("⚠  " + texto)
    run.font.size = Pt(11)
    run.font.color.rgb = AMBAR
    run.bold = True


def tabla(doc, encabezados, filas):
    """Tabla simple con cabecera en negrita."""
    t = doc.add_table(rows=len(filas) + 1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, enc in enumerate(encabezados):
        hdr[i].text = enc
    for c in hdr:
        for pr in c.paragraphs:
            for run in pr.runs:
                run.bold = True
    for i, fila in enumerate(filas, 1):
        for j, val in enumerate(fila):
            t.rows[i].cells[j].text = val
    return t


doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


# ═══ PORTADA ═══
titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = titulo.add_run("Sami")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = VERDE

subt = doc.add_paragraph()
subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subt.add_run(
    "Diario digital de acompañamiento emocional\n"
    "para estudiantes de secundaria (12 a 17 años)"
)
r.font.size = Pt(13)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Documento de defensa del proyecto.\n"
    "Lenguaje claro para revisión clínica y académica."
)
r.italic = True
r.font.color.rgb = GRIS_MEDIO


# ═══ 0. CÓMO LEER ESTE DOCUMENTO ═══
doc.add_paragraph()
h1(doc, "Cómo leer este documento")

parrafo(
    doc,
    "Este documento describe el proyecto Sami sin tecnicismos. Está pensado "
    "para que cualquier persona —el asesor de tesis, la psicóloga del "
    "colegio, un padre de familia, el jurado de sustentación— pueda leerlo "
    "de corrido y entender qué hace el sistema, cómo lo hace, quién lo usa "
    "y por qué cada decisión está tomada como está.",
)

parrafo(doc, "Está organizado en tres bloques:")
bullet(doc, "Bloque 1 (secciones 1 a 5): qué es Sami, quién lo usa y cómo entra en el día a día del colegio.")
bullet(doc, "Bloque 2 (secciones 6 a 11): cómo funciona por dentro — los ciclos de 14 días, el análisis del texto, la evaluación clínica, el reporte para la psicóloga, los protocolos de crisis.")
bullet(doc, "Bloque 3 (secciones 12 a 14): un ejemplo de reporte real, lo que Sami NO hace, y el plan de implementación.")

parrafo(
    doc,
    "Los apéndices al final listan, palabra por palabra, los ítems del "
    "PHQ-A y GAD-7 con el lenguaje juvenil peruano que el sistema reconoce. "
    "Sirven como anexo técnico de revisión clínica.",
)


# ═══ 1. QUÉ ES SAMI ═══
h1(doc, "1. Qué es Sami")

parrafo(
    doc,
    "Sami es un diario digital con cerebro clínico. El alumno escribe cómo "
    "se siente durante 14 días seguidos en su propio lenguaje, y el sistema "
    "lee ese texto para detectar señales tempranas de depresión, ansiedad, "
    "estrés, problemas de sueño, soledad o crisis. Al cabo de los 14 días, "
    "la psicóloga del colegio recibe un reporte con puntajes clínicos "
    "oficiales (PHQ-A, GAD-7), los criterios DSM-5 activados y las frases "
    "exactas del alumno que dispararon cada señal.",
)

parrafo(doc, "")
parrafo(
    doc,
    "Sami mide. La psicóloga diagnostica.",
    negrita=True,
)
parrafo(
    doc,
    "Esta es la línea ética y legal del proyecto. Sami nunca le dice al "
    "alumno que tiene depresión, ni recomienda medicación, ni reemplaza a "
    "la profesional. Lo que hace es entregarle a la psicóloga una foto "
    "clara y estructurada de las últimas dos semanas, para que ella llegue "
    "a la sesión con contexto real y tome decisiones clínicas con su "
    "criterio.",
)


# ═══ 2. EL SISTEMA EN UN MAPA ═══
h1(doc, "2. El sistema en un mapa")

parrafo(
    doc,
    "Vista general en una página. Cada parte se desarrolla en las "
    "secciones siguientes.",
    italic=True,
)

mapa = doc.add_paragraph()
r = mapa.add_run(
    "FUNDAMENTO CLÍNICO\n"
    "──────────────────\n"
    "DSM-5 (criterios diagnósticos oficiales)\n"
    "  ├─ Trastorno Depresivo Mayor → PHQ-A (9 ítems)\n"
    "  └─ Trastorno de Ansiedad Generalizada → GAD-7 (7 ítems)\n\n"
    "EL ALUMNO ESCRIBE (14 días por ciclo)\n"
    "──────────────────\n"
    "Día 1 a Día 14:\n"
    "  Diario libre. Escribe cuando quiere, en sus palabras.\n"
    "  El sistema lee cada entrada al instante y detecta\n"
    "  señales de los 16 ítems clínicos (9 PHQ-A + 7 GAD-7).\n\n"
    "EL SISTEMA EVALÚA\n"
    "──────────────────\n"
    "Para cada ítem cuenta: ¿en cuántos días apareció?\n"
    "  0 días     → 0 puntos  (\"Nunca\")\n"
    "  1–7 días   → 1 punto   (\"Algunos días\")\n"
    "  8–11 días  → 2 puntos  (\"Más de la mitad de los días\")\n"
    "  12–14 días → 3 puntos  (\"Casi todos los días\")\n\n"
    "Suma todos los ítems:\n"
    "  PHQ-A total: 0–27  → Mínima / Leve / Moderada / Severa\n"
    "  GAD-7 total: 0–21  → Mínima / Leve / Moderada / Severa\n\n"
    "ENTREGA A LA PSICÓLOGA (día 14)\n"
    "──────────────────\n"
    "NO dice \"tiene depresión\".\n"
    "Dice, por ejemplo:\n"
    "  • Puntaje PHQ-A: 12/27 → Moderada\n"
    "  • Criterios DSM-5 con presencia: A1, A2, A3, A4, A6, A7\n"
    "  • Eventos críticos: día 8 (\"no sirvo para esto\")\n"
    "  • Confiabilidad: 71% (10/14 días escritos)\n\n"
    "LA PSICÓLOGA EN LA SESIÓN\n"
    "──────────────────\n"
    "Confirma, ajusta o descarta lo que el sistema observó.\n"
    "Diagnostica si corresponde (eso lo hace ella, no Sami).\n"
    "El diario sigue corriendo en paralelo: Ciclo 2 ya arrancó."
)
r.font.name = "Consolas"
r.font.size = Pt(9)


# ═══ 3. QUIÉNES USAN SAMI ═══
h1(doc, "3. Quiénes usan Sami")

parrafo(
    doc,
    "El sistema tiene cuatro tipos de usuario, cada uno con una vista "
    "diferente y permisos diferentes.",
)
parrafo(doc, "")

tabla(
    doc,
    ["Usuario", "Qué hace y qué ve"],
    [
        (
            "Alumno (12 a 17 años)",
            "Escribe su diario, recibe consejos personalizados de Sami después "
            "de cada entrada, ve su racha de días, su calendario de ciclos, "
            "sus citas con la psicóloga, mensajes de ella, recursos de ayuda "
            "y líneas de emergencia. NUNCA ve puntajes ni criterios DSM-5: el "
            "sistema lo protege de auto-etiquetarse.",
        ),
        (
            "Psicóloga del colegio",
            "Ve el dashboard de toda la cohorte de alumnos con su estado "
            "actual. Recibe alertas críticas en tiempo real (crisis suicida, "
            "bullying, violencia familiar). Al cierre de cada ciclo de 14 días "
            "ve un reporte con PHQ-A, GAD-7, criterios DSM-5 activados, "
            "eventos del ciclo y frases concretas. Agenda citas. Deja notas "
            "clínicas privadas. Manda mensajes al alumno.",
        ),
        (
            "Tutor del aula",
            "Recibe solo alertas críticas (crisis suicida, bullying detectado, "
            "violencia familiar) para activar el protocolo escolar. No accede "
            "al detalle clínico ni al diario completo.",
        ),
        (
            "Administrador del sistema",
            "Configura los prompts del diario, calibra los umbrales del modelo "
            "de IA, audita los accesos al sistema, gestiona los respaldos de "
            "la base de datos. Garantiza el cumplimiento de la Ley 29733 "
            "(protección de datos personales).",
        ),
        (
            "Padres o apoderado",
            "No acceden al sistema. Reciben notificación solo en casos de "
            "crisis, autolesión o derivación clínica, según el protocolo "
            "firmado en el consentimiento informado.",
        ),
    ],
)


# ═══ 4. CUÁNDO Y CÓMO ENTRA SAMI ═══
h1(doc, "4. Cuándo y cómo entra Sami en el colegio")

h3(doc, "Modelo de despliegue: universal")

parrafo(
    doc,
    "Sami se entrega como diario institucional a todos los alumnos de "
    "secundaria del colegio. No es un servicio voluntario ni una "
    "derivación clínica: es una herramienta de detección temprana que "
    "acompaña a toda la cohorte.",
)

parrafo(
    doc,
    "¿Por qué universal y no por derivación? Porque los adolescentes en "
    "mayor riesgo suelen ser los que menos piden ayuda. Si Sami solo "
    "entrara cuando un tutor detecta un problema, llegaríamos tarde a los "
    "casos que más nos importa atrapar a tiempo. Que todos lo usen "
    "elimina el sesgo y normaliza el acto de escribir sobre lo que uno "
    "siente.",
)

h3(doc, "Onboarding institucional")

parrafo(doc, "Al inicio del año escolar (o cuando un alumno nuevo ingresa):")
bullet(doc, "Charla en el aula con la psicóloga del colegio: qué es Sami, qué pasa con lo que escriben, qué se notifica y a quién.")
bullet(doc, "Entrega del consentimiento informado para los padres (físico o digital).")
bullet(doc, "El alumno firma el asentimiento adaptado a su edad la primera vez que entra al sistema.")
bullet(doc, "Recibe credenciales personales con clave inicial que cambia en el primer ingreso.")

h3(doc, "Consentimiento informado de los padres + asentimiento del alumno")

parrafo(
    doc,
    "Por tratarse de menores de edad, el uso del sistema requiere DOS "
    "firmas previas, exigidas por el Código de los Niños y Adolescentes "
    "(Ley 27337) y la Ley de Salud Mental (Ley 30947):",
)

bullet(doc, "Consentimiento del padre, madre o apoderado legal. Documento físico/digital firmado al inicio del año.")
bullet(doc, "Asentimiento del alumno, en versión adaptada para su edad. Se firma digitalmente la primera vez que ingresa.")

parrafo(
    doc,
    "Ambos documentos explican en lenguaje claro: qué datos se recolectan, "
    "quién los ve, en qué situaciones se comparten con la familia (crisis, "
    "ideación suicida, bullying, autolesión, violencia familiar), cuánto "
    "tiempo se almacenan, y cómo el alumno o sus padres pueden solicitar "
    "el borrado de la información (Ley 29733).",
)

alerta(
    doc,
    "La confidencialidad NO es absoluta. Por ser menores de edad, ante "
    "una crisis o riesgo de daño, el sistema notifica a tutor y padres "
    "según protocolo. Esto se informa con anticipación en el consentimiento."
)


# ═══ 5. EL DÍA A DÍA DEL ALUMNO ═══
h1(doc, "5. El diario, día a día, desde la perspectiva del alumno")

parrafo(
    doc,
    "Para el alumno, Sami no es un cuestionario ni una evaluación. Es un "
    "diario privado al que puede entrar cuando quiera. Esta es su "
    "experiencia típica:",
)

bullet(doc, "Entra a la aplicación con su usuario y clave.")
bullet(doc, "Ve su menú principal: \"Mi diario\", \"Mi historial\", \"Recursos de ayuda\", \"Mi perfil\".")
bullet(doc, "Al abrir el diario, Sami lo saluda y le propone un prompt sencillo: \"¿Cómo amaneciste hoy?\", \"¿Qué fue lo más fuerte de tu día?\", \"¿Hay algo que te esté preocupando?\".")
bullet(doc, "El alumno escribe lo que quiere, en sus palabras. No hay mínimo ni máximo.")
bullet(doc, "Al terminar la entrada, Sami le devuelve un consejo personalizado según lo que escribió (sueño, respiración, hablar con alguien, hacer ejercicio, etc.).")
bullet(doc, "Ve su racha (\"Llevas 6 días escribiendo este ciclo\") y dónde va del ciclo de 14 días.")
bullet(doc, "Puede consultar la pestaña de recursos: líneas de ayuda, contacto de Bienestar UPC, consejos de autocuidado.")

parrafo(
    doc,
    "El alumno no tiene que escribir todos los días. La confiabilidad del "
    "reporte final depende de cuántos días escribió, pero el sistema no "
    "lo presiona. Lo único que hace es recordarle suavemente si lleva "
    "varios días sin entrar.",
)


# ═══ 6. POR QUÉ 14 DÍAS ═══
h1(doc, "6. Por qué los ciclos duran 14 días")

parrafo(
    doc,
    "Porque el cuestionario clínico que usamos como base, el PHQ-A "
    "(Johnson, 2002), pregunta literalmente: \"en las últimas 2 semanas, "
    "¿cuántos días ha sentido esto?\". Trabajar en bloques de 14 días no "
    "es una elección arbitraria — es respetar la ventana temporal que el "
    "instrumento clínico ya definió.",
)

parrafo(
    doc,
    "Lo mismo aplica al GAD-7 (Spitzer, 2006) para ansiedad: también "
    "pregunta sobre los últimos 14 días. Por eso ambos cuestionarios "
    "comparten ventana, y por eso un ciclo de Sami siempre cubre "
    "exactamente ese período.",
)


# ═══ 7. EL CICLO DE 14 DÍAS ═══
h1(doc, "7. Cómo arranca, cierra y se renueva un ciclo")

h3(doc, "Cómo arranca el primer ciclo")

parrafo(
    doc,
    "El Día 1 del primer ciclo es el día en que el alumno escribe su "
    "primera entrada. NO es una fecha del calendario impuesta por el "
    "colegio. Si Juan escribe por primera vez el 1 de mayo, ese es su "
    "Día 1. Si María escribe por primera vez el 8 de mayo, ese es su "
    "Día 1. Cada alumno tiene su propio reloj.",
)

h3(doc, "Qué pasa durante los 14 días")

bullet(doc, "El alumno escribe cuando quiere. No es obligatorio escribir todos los días.")
bullet(doc, "Cada entrada se analiza al instante (ver sección 8).")
bullet(doc, "Si en cualquier momento aparece señal de crisis, el sistema actúa de inmediato sin esperar al final del ciclo (ver sección 11).")

h3(doc, "Qué pasa el día 14")

parrafo(
    doc,
    "El sistema cierra el ciclo automáticamente y calcula los puntajes "
    "PHQ-A y GAD-7 (ver sección 9). De ahí salen dos caminos:",
)

tabla(
    doc,
    ["Resultado", "Qué hace el sistema"],
    [
        (
            "Bajo (Mínima o Leve)",
            "El reporte se archiva. La psicóloga no necesita intervenir. "
            "El alumno arranca su Ciclo 2 al día siguiente.",
        ),
        (
            "Medio o Alto (Moderada / Severa)",
            "El reporte llega al dashboard de la psicóloga marcado como "
            "\"derivación sugerida\". Ella agenda una cita con el alumno.",
        ),
    ],
)

h3(doc, "Qué pasa al día siguiente del cierre")

parrafo(
    doc,
    "Sea cual sea el resultado del ciclo anterior, el Día 15 arranca un "
    "Ciclo 2 automáticamente. El diario no se detiene. Si el alumno "
    "escribe ese día, esa entrada es el Día 1 de su Ciclo 2.",
    negrita=True,
)

parrafo(
    doc,
    "Esto es clave: el diario sigue corriendo aunque la psicóloga tenga "
    "que ver al alumno. Eso lo cuida (no se queda sin su diario en un "
    "momento delicado) y le da a ella información fresca cuando se "
    "sientan a conversar.",
)

h3(doc, "La cita con la psicóloga, cuando aplica")

parrafo(
    doc,
    "Si el reporte salió Moderada o Severa, la psicóloga programa una "
    "cita en los días siguientes — por ejemplo, el día 19. Esa cita pasa "
    "EN PARALELO al Ciclo 2 que ya está corriendo, no lo interrumpe.",
)

parrafo(
    doc,
    "Cuando la psicóloga se sienta con el alumno el día 19, ve en su "
    "pantalla dos cosas:",
)

bullet(doc, "El reporte completo del Ciclo 1 (los 14 días que dispararon la alerta). Eso es lo que conversan.")
bullet(doc, "Las entradas de los días 1 al 5 del Ciclo 2 que ya está corriendo, como contexto reciente. Le permite preguntar \"¿cómo te fue esta última semana después de aquello?\".")

parrafo(
    doc,
    "Hablan, ella deja sus notas privadas, marca la cita como completada. "
    "El Ciclo 2 sigue corriendo sin que nada cambie. Cierra normalmente "
    "al cumplir sus 14 días.",
)

h3(doc, "Resumen visual")

mapa = doc.add_paragraph()
r = mapa.add_run(
    "                                          cita (día 19)\n"
    "                                              │\n"
    "                                              ▼\n"
    "   Ciclo 1                                Ciclo 2\n"
    "   ┌────────────────┐                     ┌────────────────┐\n"
    "   │ día 1 .... 14  │                     │ día 1 .... 14  │\n"
    "   └────────────────┘                     └────────────────┘\n"
    "       ↑                                       ↑\n"
    "       cierre por tiempo,             ya arrancó al día siguiente.\n"
    "       reporte PHQ-A Moderada,        los primeros 5 días son\n"
    "       va al dashboard de la          contexto reciente para la\n"
    "       psicóloga                      sesión del día 19"
)
r.font.name = "Consolas"
r.font.size = Pt(9)

parrafo(doc, "")
parrafo(
    doc,
    "En una frase: el ciclo dura 14 días, cierra solo, al día siguiente "
    "empieza el ciclo nuevo, y la cita con la psicóloga (cuando aplica) "
    "ocurre en paralelo sin detener el diario. La única excepción a esta "
    "regla es la crisis, que sí adelanta el cierre (ver sección 11).",
    italic=True,
)


# ═══ 8. CÓMO LEE SAMI EL TEXTO ═══
h1(doc, "8. Cómo \"lee\" Sami lo que escribe el alumno")

parrafo(
    doc,
    "El alumno escribe en su lenguaje natural — frases sueltas, "
    "modismos, jerga juvenil peruana. Sami no le pide que conteste "
    "preguntas con números ni con opciones múltiples. Para entender ese "
    "texto libre, el sistema combina dos herramientas que trabajan en "
    "paralelo.",
)

h3(doc, "Herramienta 1: lista de palabras clínicas adaptadas")

parrafo(
    doc,
    "Cada uno de los 16 ítems del PHQ-A y GAD-7 tiene asociada una lista "
    "de palabras o frases que, si aparecen en el diario, activan ese ítem. "
    "Estas listas se basan en la adaptación del PHQ-9 al español "
    "(Diez-Quevedo, 2001) y en revisión del lenguaje juvenil local.",
)

parrafo(doc, "Por ejemplo:")
bullet(doc, "El alumno escribe: \"estoy triste\" → activa el ítem 2 del PHQ-A (estado de ánimo deprimido).")
bullet(doc, "El alumno escribe: \"no puedo dormir\" → activa el ítem 3 del PHQ-A (alteración del sueño).")
bullet(doc, "El alumno escribe: \"se burlan de mí\" → activa el protocolo de bullying (Ley 29719).")
bullet(doc, "El alumno escribe: \"no quiero vivir\" → activa el ítem 9 del PHQ-A y dispara crisis inmediata.")

parrafo(
    doc,
    "La lista completa de palabras adaptadas, ítem por ítem, está en los "
    "apéndices A, B y C de este documento.",
)

h3(doc, "Herramienta 2: modelo de inteligencia artificial BERT")

parrafo(
    doc,
    "Las palabras exactas no siempre alcanzan. Un adolescente puede "
    "expresar tristeza sin usar la palabra \"triste\" — puede escribir "
    "\"ya no me motiva nada\", \"siento que todo está oscuro\", \"no le "
    "encuentro sentido\". Para captar ese sentido más allá de las "
    "palabras literales, Sami usa BERT.",
)

parrafo(
    doc,
    "BERT es un modelo de inteligencia artificial entrenado para entender "
    "el español. El modelo concreto que usa Sami se llama BETO + XNLI: "
    "fue entrenado por la Universidad de Chile específicamente en español "
    "(Cañete, 2020). Es público, abierto y se ejecuta localmente en el "
    "servidor del colegio — la información del alumno NO se envía a "
    "ninguna API externa ni a OpenAI ni a ningún tercero.",
)

parrafo(doc, "Lo que hace BERT, en términos sencillos:")
bullet(doc, "Lee la entrada del diario como un ser humano la leería.")
bullet(doc, "Para cada uno de los 16 ítems, calcula un puntaje de \"qué tan probable es que el alumno esté expresando este síntoma\", aunque no use las palabras exactas.")
bullet(doc, "Si la palabra clínica también aparece (herramienta 1), refuerza el puntaje. Si no aparece, el puntaje sigue siendo válido pero con menor confianza.")

parrafo(
    doc,
    "Cuando las dos herramientas coinciden, la señal es muy confiable. "
    "Cuando solo una encuentra evidencia, el sistema lo marca con menor "
    "confianza para que la psicóloga lo considere.",
)

h3(doc, "El papel del DSM-5")

parrafo(
    doc,
    "El DSM-5 (Manual Diagnóstico y Estadístico de Trastornos Mentales, "
    "American Psychiatric Association, 2013) es la referencia clínica "
    "oficial. Define qué es exactamente un trastorno depresivo mayor o "
    "un trastorno de ansiedad generalizada, qué criterios deben cumplirse, "
    "durante cuánto tiempo, etc.",
)

parrafo(
    doc,
    "Sami no diagnostica trastornos del DSM-5. Lo que hace es etiquetar "
    "cada ítem del PHQ-A y GAD-7 con el criterio DSM-5 correspondiente "
    "(por ejemplo: \"ítem 2 del PHQ-A = criterio A1 del DSM-5 para "
    "depresión\"). Eso le permite a la psicóloga, al ver el reporte, "
    "saber rápidamente qué criterios diagnósticos tienen presencia y "
    "cuáles no, sin tener que cruzar tablas manualmente.",
)


# ═══ 9. CÓMO EVALÚA AL FINAL DEL CICLO ═══
h1(doc, "9. Cómo evalúa al final de los 14 días")

parrafo(
    doc,
    "Al cumplirse el día 14, el sistema usa la misma lógica que usaría "
    "la psicóloga al pasar el cuestionario en papel: para cada ítem cuenta "
    "en cuántos días distintos apareció el síntoma y traduce ese número "
    "a uno de los 4 niveles oficiales del PHQ-A:",
    negrita=True,
)
parrafo(doc, "")

tabla(
    doc,
    ["Días con el síntoma", "Puntos del ítem"],
    [
        ("0 días", "0 — Nunca"),
        ("1 – 7 días", "1 — Algunos días"),
        ("8 – 11 días", "2 — Más de la mitad de los días"),
        ("12 – 14 días", "3 — Casi todos los días"),
    ],
)

h3(doc, "Por qué la tabla se ve así")

parrafo(
    doc,
    "La tabla no la inventamos por nuestra cuenta. Es la traducción "
    "matemática literal de las cuatro frases que el autor original del "
    "cuestionario, Johnson (2002), escribió para que el paciente "
    "respondiera el PHQ-A. Las frases son las que aparecen en la "
    "columna de la derecha: \"Nunca\", \"Algunos días\", \"Más de la "
    "mitad de los días\" y \"Casi todos los días\". Si el período de "
    "evaluación son 14 días, esas palabras tienen un significado "
    "aritmético exacto:",
)

bullet(doc, "\"Nunca\" significa, literalmente, 0 días.")
bullet(doc, "\"Algunos días\" cubre entre 1 y 7 días — apareció, pero no es la mayoría.")
bullet(doc, "\"Más de la mitad de los días\" significa, por definición, más de la mitad de 14 — es decir, de 8 a 11.")
bullet(doc, "\"Casi todos los días\" cubre los 12 a 14 días restantes.")

parrafo(
    doc,
    "Ejemplo cotidiano: si alguien dice \"te vi enojado más de la mitad "
    "de los días esta semana\", todos entendemos que son al menos 4 de 7 "
    "días. No hay que ser psicólogo para llegar a ese número, viene de "
    "las palabras. Con el PHQ-A sobre 14 días pasa lo mismo: \"más de la "
    "mitad\" tiene que ser estrictamente más de 7, o sea 8 o más.",
)

parrafo(
    doc,
    "En otras palabras: el sistema no inventa un puntaje. Solo cuenta "
    "cuántos días apareció el síntoma y lo traduce a la frase que el "
    "cuestionario oficial usaría. Esa frase ya tiene su puntaje fijado "
    "por el PHQ-A.",
    negrita=True,
)

h3(doc, "Suma final y niveles oficiales")

parrafo(
    doc,
    "El sistema suma los puntos de los 9 ítems del PHQ-A (rango 0-27) y "
    "los 7 ítems del GAD-7 (rango 0-21), y aplica los puntos de corte "
    "oficiales de cada instrumento:",
)
parrafo(doc, "")

parrafo(doc, "PHQ-A — depresión", negrita=True)
tabla(
    doc,
    ["Puntaje", "Nivel"],
    [
        ("0 – 4", "Mínima"),
        ("5 – 9", "Leve"),
        ("10 – 14", "Moderada — derivación clínica"),
        ("15 – 19", "Moderadamente severa — urgente"),
        ("20 – 27", "Severa — emergencia"),
    ],
)

parrafo(doc, "")
parrafo(doc, "GAD-7 — ansiedad", negrita=True)
tabla(
    doc,
    ["Puntaje", "Nivel"],
    [
        ("0 – 4", "Mínima"),
        ("5 – 9", "Leve"),
        ("10 – 14", "Moderada"),
        ("15 – 21", "Severa"),
    ],
)


# ═══ 10. QUÉ VE LA PSICÓLOGA ═══
h1(doc, "10. Qué ve la psicóloga en su panel")

parrafo(
    doc,
    "La psicóloga del colegio entra al sistema con su usuario propio y "
    "accede a un panel diseñado para que pueda priorizar su tiempo y "
    "actuar primero donde más se necesita. El panel está dividido en "
    "tres bandejas, ordenadas por urgencia.",
)

h3(doc, "Bandeja 1 — Alertas críticas (rojo)")

parrafo(
    doc,
    "Lo que pasó hoy y requiere acción inmediata: crisis suicida "
    "detectada, bullying explícito, violencia familiar. Cada alerta "
    "muestra el alumno, el día y la frase exacta que la disparó. La "
    "psicóloga puede ir directo a contactar al alumno o al tutor.",
)

h3(doc, "Bandeja 2 — Cierres de ciclo pendientes de revisar (amarillo)")

parrafo(
    doc,
    "Alumnos que cerraron su ciclo de 14 días en los últimos días con "
    "resultado Moderada o Severa. Ordenados por severidad. La psicóloga "
    "revisa el reporte de cada uno y decide si agenda cita, deriva, o "
    "manda mensaje al alumno.",
)

h3(doc, "Bandeja 3 — Cohorte general (gris)")

parrafo(
    doc,
    "Todos los alumnos del colegio con su estado actual: nivel del último "
    "ciclo cerrado, en qué día del ciclo en curso van, racha de días "
    "escritos. Permite filtrar por grado/sección, por nivel de riesgo o "
    "por confiabilidad del último reporte.",
)

h3(doc, "Lo que la psicóloga puede hacer desde el panel")

bullet(doc, "Ver el reporte clínico completo de cualquier alumno: gráfico de evolución entre ciclos, PHQ-A y GAD-7 ciclo a ciclo, criterios DSM-5 activados, eventos críticos, frases que dispararon cada señal.")
bullet(doc, "Agendar citas (presenciales o virtuales) con fecha, hora, modalidad y notas.")
bullet(doc, "Marcar citas como confirmadas, completadas o canceladas.")
bullet(doc, "Dejar notas clínicas privadas que solo ella ve.")
bullet(doc, "Enviar mensajes al alumno desde su ficha.")
bullet(doc, "Exportar el reporte de un alumno a PDF para su expediente clínico.")

h3(doc, "Filosofía del panel")

parrafo(
    doc,
    "La psicóloga toma todas las decisiones clínicas. El sistema solo le "
    "entrega los parámetros calculados y las frases concretas del alumno. "
    "Si ella considera que el sistema sobre-estimó una señal, la descarta. "
    "Si considera que algo merece más atención de la que el puntaje "
    "sugiere, lo escala. El sistema nunca diagnostica, nunca medica, "
    "nunca decide por ella. Sami es su herramienta de monitoreo, no su "
    "reemplazo.",
    negrita=True,
)


# ═══ 11. PROTOCOLO DE CRISIS ═══
h1(doc, "11. Qué pasa si hay una crisis")

alerta(
    doc,
    "Por ser menores de edad, la confidencialidad NO es absoluta. "
    "El sistema escala según la ley peruana."
)

h3(doc, "Crisis suicida (cualquier día del ciclo)")

parrafo(
    doc,
    "Si el alumno escribe algo que activa el ítem 9 del PHQ-A "
    "(pensamientos de muerte, deseos de no estar, autolesión) o si BERT "
    "detecta señales de riesgo suicida, el sistema actúa de inmediato:",
)
bullet(doc, "Genera alerta crítica inmediata para la psicóloga del colegio.")
bullet(doc, "Notifica también al tutor del aula.")
bullet(doc, "Inicia comunicación con padres según el protocolo del colegio.")
bullet(doc, "Muestra al alumno las líneas de ayuda: 113 (MINSA), 100 (MIMP), 106 (SAMU).")
bullet(doc, "El ciclo en curso se cierra ese día con reporte parcial. Al día siguiente arranca un ciclo nuevo, fresco.")

h3(doc, "Bullying o acoso escolar")

parrafo(
    doc,
    "Si el sistema detecta lenguaje compatible (\"se burlan de mí\", "
    "\"me hostigan\", \"me molestan en clase\", etc.), aplica el protocolo "
    "de la Ley 29719 contra el acoso escolar y notifica a la Coordinación "
    "de Convivencia. La psicóloga lo ve marcado en su bandeja de alertas.",
)

h3(doc, "Violencia familiar")

parrafo(
    doc,
    "Si detecta frases que sugieren violencia en el hogar (\"me pegan\", "
    "\"discutimos en casa\", \"mi papá me grita\", etc.), notifica a la "
    "psicóloga y, según el caso, activa coordinación con la Línea 100 "
    "del MIMP.",
)

h3(doc, "Episodio agudo (varios síntomas el mismo día)")

parrafo(
    doc,
    "Si en una sola entrada se activan varios ítems con intensidad alta, "
    "el sistema marca ese día como \"episodio agudo\" y avisa a la "
    "psicóloga, aunque no sea crisis suicida. Ella puede decidir adelantar "
    "la sesión.",
)


# ═══ 12. EJEMPLO DE REPORTE ═══
h1(doc, "12. Ejemplo del reporte que recibe la psicóloga")

p = doc.add_paragraph()
r = p.add_run(
    "REPORTE DE CICLO #1 — Estudiante 5to A\n"
    "12 al 25 de mayo de 2026 (14 días)\n"
    "─────────────────────────────────────────\n\n"
    "Cobertura del análisis: 10/14 días escritos (71%) — Media\n\n"
    "PHQ-A inferido: 12/27 → Moderada (Johnson 2002)\n"
    "   Criterios DSM-5 sugeridos: A1, A2, A3, A4, A6, A7\n\n"
    "GAD-7 inferido: 4/21 → Mínima\n\n"
    "⚠ Alertas durante el ciclo:\n"
    "   • Día 8 — Autovaloración negativa.\n"
    "     Frase del alumno: \"no sirvo para esto\"\n\n"
    "Señales protectoras:\n"
    "   • Ejercicio (3 días) · Apoyo social (4 días)\n\n"
    "Recomendación del sistema: Evaluación clínica indicada\n"
    "─────────────────────────────────────────"
)
r.font.name = "Consolas"
r.font.size = Pt(9)

parrafo(doc, "")
parrafo(
    doc,
    "Lectura del reporte: el alumno mostró señales de depresión moderada "
    "durante el ciclo. Seis criterios del DSM-5 para depresión tienen "
    "presencia. La frase \"no sirvo para esto\" del día 8 fue un "
    "indicador particularmente claro de autovaloración negativa. Por otro "
    "lado, el alumno también mostró señales protectoras: hizo ejercicio "
    "3 días y mencionó apoyo social en 4. La psicóloga, con esta foto, "
    "decide si agenda cita, qué temas conversar, y si activa derivación "
    "externa.",
    italic=True,
)


# ═══ 13. LO QUE SAMI NO HACE ═══
h1(doc, "13. Lo que Sami NO hace (y es a propósito)")

tabla(
    doc,
    ["Lo que NO hace", "Lo que SÍ hace"],
    [
        (
            "Decirle al alumno que tiene depresión.",
            "Marcar al ítem como activo en el reporte para la psicóloga.",
        ),
        (
            "Recomendar medicación.",
            "Sugerir derivación a evaluación clínica.",
        ),
        (
            "Reemplazar a la psicóloga.",
            "Darle a la psicóloga contexto estructurado de 14 días.",
        ),
        (
            "Mostrarle puntajes al alumno.",
            "Mostrarle consejos personalizados y recursos de ayuda.",
        ),
        (
            "Garantizar confidencialidad absoluta.",
            "Notificar a tutor y padres en casos de crisis, según protocolo informado en el consentimiento.",
        ),
        (
            "Enviar datos a APIs externas (OpenAI, Google, etc.).",
            "Procesar todo en el servidor local del colegio con un modelo abierto.",
        ),
    ],
)


# ═══ 14. PLAN DE IMPLEMENTACIÓN ═══
h1(doc, "14. Plan de implementación")

h3(doc, "Stack tecnológico")

parrafo(doc, "Sami se construye con tecnologías libres y de uso clínico estable:")
bullet(doc, "Backend: FastAPI (Python) — manejo de la API, la lógica del diario y la conexión con el modelo de IA.")
bullet(doc, "Base de datos: SQLite — almacena alumnos, entradas del diario, ciclos, citas, notas clínicas, logs de auditoría.")
bullet(doc, "Modelo de IA: BETO + XNLI (Recognai/bert-base-spanish-wwm-cased-xnli) — modelo público de la Universidad de Chile, ejecutado localmente.")
bullet(doc, "Frontend: Vue 3 + Vite + Tailwind — interfaz web responsive para alumno, psicóloga y admin.")
bullet(doc, "Autenticación: JWT con tres roles (estudiante, psicólogo, admin).")
bullet(doc, "Gráficos: Chart.js para evolución del riesgo entre ciclos.")

h3(doc, "Fases del proyecto")

parrafo(doc, "El proyecto se construyó por sprints (períodos de 2 a 5 semanas):")

tabla(
    doc,
    ["Sprint", "Objetivo", "Estado"],
    [
        ("1", "Autenticación, consentimiento, perfil del alumno, panel admin básico.", "Hecho"),
        ("2", "Panel de la psicóloga con lista de alumnos y último estado.", "Hecho"),
        ("3", "Motor de análisis NLP con BERT y mapeo a PHQ-A, GAD-7, DSM-5.", "Hecho"),
        ("4", "Recomendaciones, recursos UPC, historial emocional del alumno, notificaciones.", "Hecho"),
        ("5", "Dashboard de monitoreo, alertas tempranas, agenda de citas, notas clínicas.", "Hecho"),
        ("6", "Panel admin completo: config del sistema, auditoría, respaldos, calibración del modelo.", "Hecho"),
        ("7", "Módulo de diario (entrada libre diaria, ciclos de 14 días, análisis por entrada).", "En curso"),
        ("8", "Refinamientos de UX, filtros del psicólogo, PDF clínico, monitoreo en tiempo real.", "Pendiente"),
    ],
)

h3(doc, "Mockups y diseño visual")

parrafo(
    doc,
    "El diseño visual del sistema se valida en Figma antes de codificarse. "
    "La identidad visual usa tres colores principales: verde agua "
    "(esperanza, salud), blanco (claridad, neutralidad) y negro (texto, "
    "seriedad clínica). Sin emojis decorativos. Tipografía Work Sans para "
    "transmitir cercanía sin perder formalidad.",
)

parrafo(doc, "Pantallas mockeadas en Figma:")
bullet(doc, "Login y registro del alumno.")
bullet(doc, "Consentimiento y asentimiento.")
bullet(doc, "Menú principal del alumno (Mi diario, Mi historial, Recursos, Mi perfil).")
bullet(doc, "Vista del diario: prompt de Sami + zona de escritura + consejo personalizado al cierre.")
bullet(doc, "Historial del alumno: calendario de ciclos cerrados + racha + entradas pasadas.")
bullet(doc, "Recursos: líneas de ayuda y consejos de autocuidado.")
bullet(doc, "Panel de la psicóloga: dashboard con las 3 bandejas + ficha por alumno + reporte de ciclo + agenda de citas + notas privadas.")
bullet(doc, "Panel del admin: configuración del sistema, auditoría, respaldos, calibración de umbrales BERT.")

h3(doc, "Garantías éticas y legales que el sistema implementa")

bullet(doc, "Consentimiento explícito de padres + asentimiento del alumno antes del primer uso.")
bullet(doc, "Procesamiento local del análisis NLP (sin envío de datos a terceros).")
bullet(doc, "Logs de auditoría: cada acceso al sistema queda registrado (Ley 29733).")
bullet(doc, "Roles y permisos diferenciados: el alumno no ve puntajes, el tutor no ve el detalle clínico, los padres solo reciben notificaciones de crisis.")
bullet(doc, "Protocolos de escalamiento documentados para crisis suicida, bullying (Ley 29719) y violencia familiar.")
bullet(doc, "Posibilidad de borrar la cuenta y los datos del alumno bajo solicitud (Ley 29733).")

h3(doc, "Lo que queda pendiente para cerrar al 100%")

bullet(doc, "Filtro por nivel de riesgo en el panel de la psicóloga (mejora de UX).")
bullet(doc, "Generación de PDF clínico programático (hoy se exporta con \"Imprimir\" del navegador).")
bullet(doc, "Monitoreo en tiempo real de logs de auditoría por WebSocket (hoy con polling cada 5s, ya funcional).")


# ═══ APÉNDICE A: PHQ-A ÍTEMS Y PALABRAS ADAPTADAS ═══

doc.add_page_break()
h1(doc, "Apéndice A — Ítems del PHQ-A y palabras adaptadas")

parrafo(
    doc,
    "El PHQ-A tiene 9 preguntas que se corresponden 1:1 con los criterios "
    "del DSM-5 para depresión. Para cada ítem el sistema detecta "
    "automáticamente las palabras o frases adaptadas al español "
    "(usadas por adolescentes peruanos) que aparecen en el diario. Estas "
    "listas se basan en la adaptación del PHQ-9 al español (Diez-Quevedo "
    "et al., 2001) y en revisión de lenguaje juvenil local.",
    italic=True,
)
parrafo(doc, "")

# Cada ítem PHQ-A con su contenido
PHQA_ITEMS = [
    (1, "Pérdida de interés o placer (anhedonia)",
     "Criterio DSM-5 A2",
     "Cuéntame, en estas últimas semanas ¿cómo has estado con las cosas que antes te gustaban —música, salir, hobbies, estudiar—? ¿Sigues disfrutándolas como antes?",
     ["no disfruto", "nada me da alegría", "no me gusta nada",
      "perdí el interés", "ya nada me importa", "sin ganas"]),
    (2, "Estado de ánimo deprimido",
     "Criterio DSM-5 A1",
     "¿Y tu ánimo cómo ha estado? Quiero saber si te has sentido decaído/a, triste o sin esperanza estos días.",
     ["triste", "deprim(ido/ida)", "vacío", "vacía",
      "sin esperanza", "desesperanza", "decaído", "decaida"]),
    (3, "Alteración del sueño",
     "Criterio DSM-5 A4",
     "Hablemos un poco de tu sueño. ¿Has estado durmiendo bien, te cuesta conciliarlo, o quizás duermes más de la cuenta?",
     ["no puedo dormir", "insomnio", "duermo mucho",
      "duermo todo el día", "me despierto", "no concilio", "trasnoché"]),
    (4, "Fatiga o pérdida de energía",
     "Criterio DSM-5 A6",
     "¿Y la energía? ¿Sientes que tienes fuerzas para tu día a día, o has estado más cansado/a de lo normal, incluso después de descansar?",
     ["cansado/cansada", "agotado/agotada", "sin energía",
      "fatiga", "no tengo fuerzas", "exhausto/exhausta"]),
    (5, "Cambio en el apetito o peso",
     "Criterio DSM-5 A3",
     "Cuéntame, ¿cómo va tu apetito? ¿Comes parecido a siempre, o has notado que comes menos —o más— de lo normal?",
     ["no tengo hambre", "como mucho", "atracón",
      "perdí el apetito", "comer de más", "sin apetito"]),
    (6, "Sentimientos de inutilidad o culpa",
     "Criterio DSM-5 A7",
     "A veces uno se siente mal consigo mismo —como si no fuera suficiente, o como si le hubiera fallado a alguien importante. ¿Te ha pasado algo así estas semanas?",
     ["inútil", "fracasado/fracasada", "soy un fracaso",
      "no sirvo", "le fallé", "culpa", "no valgo nada", "decepción"]),
    (7, "Dificultad para concentrarse",
     "Criterio DSM-5 A8",
     "¿Cómo va tu concentración? Por ejemplo, cuando lees algo, ves una serie o estás en clase, ¿logras mantenerte enfocado/a?",
     ["no me concentro", "me distraigo", "no puedo concentrarme",
      "se me va la cabeza", "no retengo", "leer y no entiendo"]),
    (8, "Enlentecimiento o agitación psicomotora",
     "Criterio DSM-5 A5",
     "Cuéntame si has notado algo distinto en cómo te mueves o hablas: ¿alguien te ha dicho que estás más lento/a de lo usual, o tú sientes que no puedes parar de moverte?",
     ["muy lento", "me muevo lento", "no puedo quedarme quieto",
      "inquieto/inquieta", "agitado/agitada", "acelerado/acelerada"]),
    (9, "Ideación suicida o de autolesión",
     "Criterio DSM-5 A9",
     "Esta pregunta la hago con cuidado, porque me importa cómo estás. ¿En estas semanas has tenido pensamientos de hacerte daño, de no estar, o de que sería mejor desaparecer? Puedes responder con confianza, esto es para cuidarte.",
     ["hacerme daño", "mejor muerto", "no quiero vivir",
      "suicidarme", "desaparecer", "no despertar",
      "terminar con todo", "autolesión", "cortarme"]),
]

for num, criterio, dsm, pregunta, kws in PHQA_ITEMS:
    h3(doc, f"PHQ-A — Ítem {num}: {criterio}")
    parrafo(doc, dsm, italic=True)
    p = doc.add_paragraph()
    r = p.add_run("Pregunta oficial adaptada: ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(f"\"{pregunta}\"")
    p = doc.add_paragraph()
    r = p.add_run("Palabras adaptadas que el sistema detecta:")
    r.bold = True
    r.font.size = Pt(11)
    for kw in kws:
        bullet(doc, f"\"{kw}\"")
    parrafo(doc, "")

# Alerta especial para ítem 9
alerta(
    doc,
    "El ítem 9 tiene umbral de detección más conservador. "
    "CUALQUIER activación dispara el protocolo de crisis inmediato."
)


# ═══ APÉNDICE B: GAD-7 ÍTEMS Y PALABRAS ADAPTADAS ═══

doc.add_page_break()
h1(doc, "Apéndice B — Ítems del GAD-7 y palabras adaptadas")

parrafo(
    doc,
    "El GAD-7 tiene 7 preguntas que evalúan los síntomas del Trastorno de "
    "Ansiedad Generalizada (DSM-5). Para cada ítem el sistema detecta "
    "automáticamente las palabras o frases adaptadas al lenguaje juvenil.",
    italic=True,
)
parrafo(doc, "")

GAD7_ITEMS = [
    (1, "Ansiedad o nerviosismo excesivo",
     "Criterio DSM-5 C1 (TAG)",
     "Cuéntame cómo has estado con los nervios. ¿Te has sentido más ansioso/a o con los nervios de punta últimamente?",
     ["nervios", "ansiedad", "tensión", "con los nervios",
      "nerviosa/nervioso", "nerviosismo"]),
    (2, "Dificultad para controlar la preocupación",
     "Criterio DSM-5 B (TAG)",
     "Y cuando empiezas a preocuparte, ¿logras parar de pensar en eso, o sientes que la preocupación se te escapa de las manos?",
     ["no puedo dejar de preocuparme", "no controlo",
      "no paro de pensar", "sobrepensar", "rumiar", "le doy vueltas"]),
    (3, "Preocupación excesiva sobre múltiples temas",
     "Criterio DSM-5 A (TAG)",
     "¿Te ha pasado que te preocupas por muchas cosas a la vez —los estudios, la familia, el futuro, la salud— y todo se junta en tu cabeza?",
     ["me preocupo por todo", "todo me preocupa",
      "preocupación constante", "preocupado/preocupada"]),
    (4, "Tensión muscular o dificultad para relajarse",
     "Criterio DSM-5 C5 (TAG)",
     "Hablemos del cuerpo: ¿logras relajarte, o sientes que estás todo el tiempo tenso/a, con el cuerpo rígido o sin poder desconectar?",
     ["no me puedo relajar", "no logro relajarme", "tensión",
      "rígido/rígida", "contracturado/contracturada"]),
    (5, "Inquietud o sensación de estar al límite",
     "Criterio DSM-5 C1 (TAG)",
     "¿Has sentido como una inquietud por dentro, como si no pudieras quedarte quieto/a, ni siquiera cuando intentas descansar?",
     ["inquieto/inquieta", "no me quedo quieto",
      "tengo que moverme", "al límite"]),
    (6, "Irritabilidad",
     "Criterio DSM-5 C4 (TAG)",
     "Cuéntame, ¿has estado más irritable o de mal humor estos días? Como si las cosas pequeñas te fastidiaran más de lo normal.",
     ["irritado/irritada", "me molesto", "me enojo rápido",
      "todo me molesta", "exploto", "malhumorado/malhumorada"]),
    (7, "Sensación de peligro o catástrofe inminente",
     "Criterio DSM-5 A (TAG) — preocupación catastrófica",
     "A veces uno tiene la sensación de que algo malo va a pasar, sin saber muy bien qué. ¿Te ha pasado algo así últimamente?",
     ["algo malo va a pasar", "siento que va a pasar algo",
      "miedo", "pánico", "catástrofe", "presentimiento"]),
]

for num, criterio, dsm, pregunta, kws in GAD7_ITEMS:
    h3(doc, f"GAD-7 — Ítem {num}: {criterio}")
    parrafo(doc, dsm, italic=True)
    p = doc.add_paragraph()
    r = p.add_run("Pregunta oficial adaptada: ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(f"\"{pregunta}\"")
    p = doc.add_paragraph()
    r = p.add_run("Palabras adaptadas que el sistema detecta:")
    r.bold = True
    r.font.size = Pt(11)
    for kw in kws:
        bullet(doc, f"\"{kw}\"")
    parrafo(doc, "")


# ═══ APÉNDICE C: KEYWORDS ADICIONALES ═══
doc.add_page_break()
h1(doc, "Apéndice C — Otros temas que detecta el sistema")

parrafo(
    doc,
    "Más allá de los 16 ítems PHQ-A/GAD-7, el sistema detecta tres temas "
    "que NO están en los cuestionarios oficiales pero son clínicamente "
    "relevantes en adolescentes y disparan protocolos específicos:",
    italic=True,
)
parrafo(doc, "")

TEMAS_EXTRA = [
    ("Bullying / acoso escolar",
     "Ley 29719 — protocolo de Convivencia Escolar",
     ["bullying", "acoso", "me molestan", "se burlan",
      "me hostigan", "matonea", "me jodén", "me discriminan"]),
    ("Conflicto familiar / violencia en el hogar",
     "Notificación al psicólogo y eventual Línea 100 (MIMP)",
     ["mi mamá", "mi papá", "mis padres", "en casa",
      "discutimos en casa", "pelea en casa", "violencia en casa",
      "mi hermano me", "mi hermana me"]),
    ("Problemas de sueño persistentes",
     "Refuerza el ítem PHQ-A #3 y se incluye en recomendaciones",
     ["no duermo", "no puedo dormir", "insomnio", "pesadillas",
      "me despierto", "duermo poco", "duermo mal"]),
]

for tema, accion, kws in TEMAS_EXTRA:
    h3(doc, tema)
    p = doc.add_paragraph()
    r = p.add_run("Acción del sistema: ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(accion)
    p = doc.add_paragraph()
    r = p.add_run("Palabras detectadas:")
    r.bold = True
    r.font.size = Pt(11)
    for kw in kws:
        bullet(doc, f"\"{kw}\"")
    parrafo(doc, "")


# ═══ BIBLIOGRAFÍA ═══
doc.add_page_break()
h1(doc, "Bibliografía mínima")

refs = [
    "American Psychiatric Association. (2013). DSM-5. APA Publishing.",
    "Johnson, J. G., Harris, E. S., Spitzer, R. L., & Williams, J. B. (2002). The Patient Health Questionnaire for Adolescents. Journal of Adolescent Health, 30(3), 196-204.",
    "Spitzer, R. L., Kroenke, K., Williams, J. B., & Löwe, B. (2006). A brief measure for assessing generalized anxiety disorder: the GAD-7. Archives of Internal Medicine, 166(10), 1092-1097.",
    "Mossman, S. A., Luft, M. J., Schroeder, H. K., et al. (2017). The GAD-7 in adolescents with generalized anxiety disorder. Annals of Clinical Psychiatry, 29(4), 227-234.",
    "Diez-Quevedo, C., et al. (2001). Validation and utility of the Patient Health Questionnaire in Spanish inpatients. Psychosomatic Medicine, 63(4), 679-686.",
    "Cañete, J., Chaperon, G., Fuentes, R., et al. (2020). Spanish Pre-Trained BERT Model and Evaluation Data. ICLR 2020.",
    "Congreso de la República del Perú. (2019). Ley 30947 — Ley de Salud Mental.",
    "Congreso de la República del Perú. (2000). Ley 27337 — Código de los Niños y Adolescentes.",
    "Congreso de la República del Perú. (2011). Ley 29719 — Convivencia sin violencia en instituciones educativas.",
    "Congreso de la República del Perú. (2011). Ley 29733 — Ley de Protección de Datos Personales.",
    "MINEDU. (2021). Resolución Viceministerial N° 169-2021. Lineamientos de convivencia escolar.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(ref)
    run.font.size = Pt(10)


doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Sami — Diario digital de acompañamiento emocional\n"
    "Tesis de pregrado, 2026"
)
r.italic = True
r.font.color.rgb = GRIS_MEDIO
r.font.size = Pt(9)


output = "/home/renzo/salud-mental-TP1/docs/Diseño_del_sistema_Sami.docx"
doc.save(output)
print(f"Documento generado: {output}")
